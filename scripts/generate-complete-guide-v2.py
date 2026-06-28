#!/usr/bin/env python3
"""
ContentForge — COMPLETE Updated Video Scripts + Instructions
Rewritten to include ALL features (20 tabs, 23 videos).
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "download", "ContentForge-Complete-Guide-v2.docx")

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h

def add_para(doc, text, bold=False, italic=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_narration(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"NARRATION: {text}")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

def add_visual(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"VISUAL: {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run.bold = True

def add_tip(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"PRO TIP: {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    run.bold = True

def add_field(doc, name, ftype, location, desc, example=None, tip=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xE8, 0x6A, 0x00)
    run = p.add_run(f"  ({ftype})")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p = doc.add_paragraph()
    run = p.add_run(f"Where: {location}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    p = doc.add_paragraph()
    run = p.add_run(f"What it does: {desc}")
    run.font.size = Pt(11)
    if example:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run("Example: ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(example)
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    if tip:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run("Tip: ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
        run = p.add_run(tip)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)

def add_video(doc, num, title, duration, desc):
    doc.add_page_break()
    add_heading(doc, f"Video {num}: {title}", level=1, color=RGBColor(0xE8, 0x6A, 0x00))
    add_para(doc, f"Duration: ~{duration} minutes", italic=True, size=10, color=RGBColor(0x66, 0x66, 0x66))
    add_para(doc, desc, size=11, space_after=12)

def build():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ============ COVER ============
    for _ in range(5): doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ContentForge")
    run.bold = True
    run.font.size = Pt(42)
    run.font.color.rgb = RGBColor(0xE8, 0x6A, 0x00)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Complete Guide: Video Scripts + Field Instructions")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("Updated edition covering all 20 tabs and 23 tutorial videos\nIncludes field-by-field instructions for every input")
    run.font.size = Pt(13)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("23 Videos | ~85 min total runtime | 20 feature tabs\nFor screen recording or AI video generation")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ============ TOC ============
    doc.add_page_break()
    add_heading(doc, "Table of Contents", level=1)
    add_para(doc, "PART 1: VIDEO SCRIPTS (23 videos)", bold=True, size=14, color=RGBColor(0xE8, 0x6A, 0x00))
    videos = [
        ("1", "Platform Overview & Welcome", "2 min"),
        ("2", "Dashboard Tour", "3 min"),
        ("3", "Uploading Videos", "5 min"),
        ("4", "Uploading Photos to Video", "4 min"),
        ("5", "Video Library Management", "3 min"),
        ("6", "AI Agent Chat", "5 min"),
        ("7", "Content Ideation Engine", "4 min"),
        ("8", "AI Generation (Thumbnails, Images, B-roll)", "5 min"),
        ("9", "Script Analyzer - Reverse-Engineer Viral Videos", "5 min"),
        ("10", "Framework Library - Reusable Templates", "3 min"),
        ("11", "Social Accounts & API Keys", "6 min"),
        ("12", "Scheduling & Calendar", "4 min"),
        ("13", "Trend Research", "3 min"),
        ("14", "Analytics Dashboard", "4 min"),
        ("15", "Insights Engine", "3 min"),
        ("16", "Auto-Comment Replies", "5 min"),
        ("17", "Voice Profile Setup", "3 min"),
        ("18", "Competitor Monitoring", "4 min"),
        ("19", "Brand Kit", "3 min"),
        ("20", "Content Repurposing", "4 min"),
        ("21", "Subtitle Translation", "3 min"),
        ("22", "Hook A/B Testing", "3 min"),
        ("23", "Multi-Format Publishing", "3 min"),
    ]
    for num, title, dur in videos:
        p = doc.add_paragraph()
        run = p.add_run(f"Video {num}: {title}")
        run.font.size = Pt(11)
        run = p.add_run(f"\t{dur}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    add_para(doc, "")
    add_para(doc, "PART 2: FIELD-BY-FIELD GUIDE (20 tabs)", bold=True, size=14, color=RGBColor(0xE8, 0x6A, 0x00))
    tabs = [
        "Dashboard", "Upload (Video + Photo + Voiceover)", "Library",
        "Ideas", "Generate", "Script Analyzer", "Framework Library",
        "Social", "Scheduled", "Calendar", "Trends", "Analytics",
        "Insights", "Comments", "Competitors", "Brand Kit", "Voice",
        "API Keys", "Assets", "Settings", "AI Agent Chat"
    ]
    for t in tabs:
        add_para(doc, f"  - {t}", size=11)

    add_para(doc, "")
    add_para(doc, "PART 3: RECORDING CHECKLIST", bold=True, size=14, color=RGBColor(0xE8, 0x6A, 0x00))
    add_para(doc, "  - Phase 1: Pre-Recording Setup", size=11)
    add_para(doc, "  - Phase 2: Recording Session", size=11)
    add_para(doc, "  - Phase 3: Post-Recording", size=11)
    add_para(doc, "  - Equipment & Software", size=11)

    # ============ INTRODUCTION ============
    doc.add_page_break()
    add_heading(doc, "Introduction", level=1)
    add_para(doc, "This document is your complete master guide to ContentForge. It contains three parts:")
    add_para(doc, "Part 1: 23 video scripts with narration, visual cues, and pro tips. Use these to record tutorial videos or generate AI videos.")
    add_para(doc, "Part 2: A field-by-field guide explaining every input, dropdown, toggle, and button across all 20 tabs. Use this as your reference manual.")
    add_para(doc, "Part 3: A recording checklist to ensure professional-quality tutorials every time.")
    add_para(doc, "")
    add_para(doc, "How to use the video scripts:", bold=True, size=12)
    add_para(doc, "1. Each video has NARRATION (what to say), VISUAL (what to show), and PRO TIP sections.")
    add_para(doc, "2. For screen recording: Follow the VISUAL cues while reading the NARRATION.")
    add_para(doc, "3. For AI video generation (Higgsfield, etc.): Use NARRATION as script input, describe VISUAL as scene directions.")
    add_para(doc, "4. Videos are self-contained. Record in any order or combine into longer compilations.")
    add_para(doc, "")
    add_para(doc, "Recording tips:", bold=True, size=12)
    add_para(doc, "- Record in 1080p or 4K so viewers can read UI text when paused")
    add_para(doc, "- Use a clean browser profile (no bookmarks bar, no extensions)")
    add_para(doc, "- Zoom into UI elements using screen magnifier")
    add_para(doc, "- Add captions/subtitles for accessibility")
    add_para(doc, "- Keep a consistent intro/outro across all videos")

    # ============ PART 1: VIDEO SCRIPTS ============
    doc.add_page_break()
    add_heading(doc, "PART 1: VIDEO SCRIPTS", level=1, color=RGBColor(0xE8, 0x6A, 0x00))

    # Video 1
    add_video(doc, "1", "Platform Overview & Welcome", "2 min",
        "A high-level introduction to ContentForge covering the complete content lifecycle.")
    add_visual(doc, "Open browser to ContentForge dashboard. Slow zoom out showing the full interface with all 20 nav tabs.")
    add_narration(doc, "Welcome to ContentForge, the all-in-one AI content automation platform that handles your entire content workflow from idea to published post. Whether you are a pet creator, fitness influencer, or business brand, ContentForge automates the repetitive work so you can focus on creating.")
    add_visual(doc, "Hover over each nav tab: Dashboard, Upload, Library, Ideas, Generate, Analyzer, Frameworks, Social, Scheduled, Calendar, Trends, Analytics, Insights, Comments, Competitors, Brand Kit, Voice, API Keys, Assets, Settings.")
    add_narration(doc, "The platform covers the complete content lifecycle across 20 tabs. You can upload videos or photos, let AI edit them with captions and voiceover, reverse-engineer viral videos to understand why they work, schedule posts across five social platforms, monitor competitors, auto-reply to comments, and learn what works through built-in analytics.")
    add_visual(doc, "Click the AI Agent floating button (bottom right). Show the chat opening.")
    add_narration(doc, "And the best part? There is a built-in AI agent with 38 tools that can do almost everything for you. Just ask it to generate ideas, analyze a viral video, schedule posts, check trends, or analyze your performance, and it executes. Let's dive into each feature.")
    add_tip(doc, "This works great as a YouTube intro or homepage hero video. Keep it under 2 minutes.")

    # Video 2
    add_video(doc, "2", "Dashboard Tour", "3 min",
        "Walk through the main dashboard, explaining each stat card and quick action.")
    add_visual(doc, "Start on the Dashboard tab. Point to each stat card.")
    add_narration(doc, "The Dashboard is your command center. At the top you see six stat cards: Total Videos, Processing, Ready, Scheduled, Published, and Average Viral Score. These update in real-time as your content moves through the pipeline.")
    add_visual(doc, "Scroll down to Quick Actions. Hover over each button.")
    add_narration(doc, "Below the stats you find Quick Actions, shortcuts to common tasks: Upload new videos, Generate content ideas with AI, Generate AI thumbnails and B-roll, View library, and Content calendar. Orange buttons are AI-powered features.")
    add_visual(doc, "Scroll to the How it works section.")
    add_narration(doc, "At the bottom there is a quick reference showing the complete workflow: Upload, transcribe, edit, score, multi-format generation, and publish. Great for new users to understand the pipeline at a glance.")
    add_visual(doc, "Point to the floating orange sparkle button bottom right.")
    add_narration(doc, "And do not forget the AI Agent. This floating button is always available no matter which tab you are on. Click it anytime to ask questions or trigger actions without navigating through menus.")
    add_tip(doc, "Mention that the niche shown in the welcome message adapts to whatever the user set in Settings.")

    # Video 3
    add_video(doc, "3", "Uploading Videos", "5 min",
        "Complete walkthrough of the video upload process, including all editing options.")
    add_visual(doc, "Click the Upload tab. Show the drag-and-drop zone.")
    add_narration(doc, "Let's upload our first video. Click the Upload tab. You can drag and drop video files here, or click to browse. ContentForge supports MP4, MOV, and WebM, and you can upload multiple files at once.")
    add_visual(doc, "Select a video file. It appears in the file list.")
    add_narration(doc, "Once selected, files appear in this list. Remove any file by clicking the X button. Now let's look at the editing options.")
    add_visual(doc, "Scroll to Auto-Editing Options. Point to each toggle.")
    add_narration(doc, "First, Burn AI Captions. This transcribes your audio and burns subtitles directly onto the video. Great for accessibility and engagement. Next, Auto-trim Silence removes dead air from your video.")
    add_visual(doc, "Show the Watermark, Background Music, Intro Clip, and Outro Clip dropdowns.")
    add_narration(doc, "You can add a watermark, background music, intro clips, and outro clips. Upload these once in the Assets tab and they appear in these dropdowns for every video.")
    add_visual(doc, "Scroll to AI Voiceover section. Toggle it on.")
    add_narration(doc, "Now here is a powerful feature. AI Voiceover. When enabled, the system generates a script based on your niche, then synthesizes it as audio using the voice you select. Choose from seven voices, set the tone, and decide whether to replace the original audio or mix the voiceover on top.")
    add_visual(doc, "Click Upload. Show the video appearing in Library with progress bar.")
    add_narration(doc, "Once you click Upload, the video enters the processing pipeline. You see it in your Library with a progress bar showing each step: probing metadata, extracting audio, transcribing with AI, analyzing viral potential, editing with FFmpeg, and generating multi-format outputs. The whole process takes about 30 to 60 seconds.")
    add_tip(doc, "The viral score is generated by AI comparing the video to current trends.")

    # Video 4
    add_video(doc, "4", "Uploading Photos to Video", "4 min",
        "How to upload still photos and have them converted into videos with voiceover, music, and captions.")
    add_visual(doc, "On the Upload tab, point to the purple Upload Photos Instead button.")
    add_narration(doc, "Did you know you can create videos from photos? ContentForge turns still images into engaging videos with Ken Burns zoom effects, AI voiceover, burned captions, and background music. Perfect for photo dumps, before-and-after content, or any time you do not have video footage.")
    add_visual(doc, "Click Upload Photos Instead. Select 3-4 photos. They appear in a purple section.")
    add_narration(doc, "Click Upload Photos Instead and select one or more images. They appear in this purple section. If you upload multiple photos, they combine into a slideshow with crossfade transitions.")
    add_visual(doc, "Show the photo settings: seconds per photo, transition, custom voiceover script.")
    add_narration(doc, "Now you have photo-specific settings. Seconds per photo controls how long each image displays, default 5 seconds. Transition controls the crossfade duration. And here is the best part, you can write a custom voiceover script, or leave it blank and let AI generate one based on your niche and current trends.")
    add_visual(doc, "Type a custom script in the voiceover script field.")
    add_narration(doc, "If you provide a script, that exact text is used for the voiceover AND the burned captions. If you leave it blank, AI writes a script tailored to your niche and brand voice.")
    add_visual(doc, "Click Upload. Show the video appearing in Library.")
    add_narration(doc, "Click Upload and the system converts your photos into a video, generates the voiceover, burns captions, adds music, and produces all three formats. The result appears in your Library looking just like a regular video, ready to publish.")
    add_tip(doc, "Show an example of a finished photo-to-video result with AI title and viral score.")

    # Video 5
    add_video(doc, "5", "Video Library Management", "3 min",
        "How to manage processed videos including publish, repurpose, translate, and delete.")
    add_visual(doc, "Navigate to Library tab. Show the grid of video cards.")
    add_narration(doc, "The Library is where all your processed videos live. Each card shows the thumbnail, viral score, status, AI-generated title, duration, file size, and hashtags.")
    add_visual(doc, "Point to the viral score circle top right of a card.")
    add_narration(doc, "This circle shows your AI-generated viral score from 0 to 100. Green means 80 or above, high viral potential. Amber means 50 to 79, decent. Red means below 50, might need improvement.")
    add_visual(doc, "Point to the status badge top left. Show different statuses.")
    add_narration(doc, "The status badge shows where the video is in the pipeline. Pending means queued. Editing means FFmpeg is processing. Ready means done and ready to publish. Published means pushed to social platforms. Failed means something went wrong, click Retry to reprocess.")
    add_visual(doc, "Point to format badges below the video preview.")
    add_narration(doc, "These green badges show which formats have been generated: 16:9 for YouTube and Facebook, 9:16 for TikTok and Reels, and 1:1 for Instagram feed. Every video automatically gets all three.")
    add_visual(doc, "Click Publish on a ready video. Show the publish dialog briefly.")
    add_narration(doc, "When a video is ready, you see action buttons. Publish opens the publishing dialog where you select platforms and schedule. Translate lets you translate the caption to 12 languages. And Repurpose appears on videos longer than 90 seconds, it uses AI to cut the video into short clips.")
    add_tip(doc, "Clips show a purple Clip badge and display their source timestamp.")

    # Video 6
    add_video(doc, "6", "AI Agent Chat", "5 min",
        "Deep dive into the AI Agent with 38 tools that can see your data and take actions.")
    add_visual(doc, "Click the floating orange sparkle button bottom right. Chat sidebar opens.")
    add_narration(doc, "This is the ContentForge AI Agent, your personal content strategist, social media manager, and analyst all in one. It is not just a chatbot. It has access to 38 tools that let it see your data and take real actions in your account.")
    add_visual(doc, "Show the welcome screen with suggested questions.")
    add_narration(doc, "When you first open it, you see suggested questions. Let's try one.")
    add_visual(doc, "Click 'What should I post this week?' Show the agent calling tools and generating a response.")
    add_narration(doc, "Watch what happens when I ask What should I post this week. The agent calls three tools in sequence: it lists your videos, reads your settings, and searches current trends. Then it synthesizes all that data into a personalized weekly strategy.")
    add_visual(doc, "Show the response with tool call indicators.")
    add_narration(doc, "These little chips show which tools the agent used. You can expand them to see the actual data returned. This transparency builds trust, you know exactly what the agent based its advice on.")
    add_visual(doc, "Type: 'Analyze this viral video: [paste URL]'")
    add_narration(doc, "The agent can also analyze viral videos. Just paste a YouTube URL and ask it to analyze. It will reverse-engineer the script, identify psychology triggers, extract the framework, and even adapt it for your niche.")
    add_visual(doc, "Show the agent confirming actions taken.")
    add_narration(doc, "After taking an action, the agent confirms what it did. You can always verify by checking the relevant tab. Conversations are saved so you can reference them later.")
    add_tip(doc, "The agent remembers context within a conversation. If you ask what is my brand handle, it already knows from previous messages.")

    # Video 7
    add_video(doc, "7", "Content Ideation Engine", "4 min",
        "How to use the Ideas tab to generate AI content ideas based on trends and analytics.")
    add_visual(doc, "Navigate to Ideas tab. Show the Generate 5 ideas button.")
    add_narration(doc, "Never know what to post? The Ideas tab uses AI to generate content ideas tailored to your niche, grounded in current trends and your actual performance data.")
    add_visual(doc, "Click Generate 5 ideas. Show loading then results.")
    add_narration(doc, "Click Generate 5 ideas and the system pulls together your current trends, analytics insights, recent videos, and niche settings. It sends all that context to the AI which generates five original content ideas. This takes about 15 to 20 seconds.")
    add_visual(doc, "Scroll through the generated ideas. Point to each element.")
    add_narration(doc, "Each idea card includes a scroll-stopping title, a predicted viral score from 0 to 100, a concept description, a full script outline with Hook, Body, and CTA sections, the recommended format, and the source which shows what data informed this idea.")
    add_visual(doc, "Expand the script outline section.")
    add_narration(doc, "The script outline is production-ready. It gives you a hook for the first three seconds, the body content, and a call to action. You can take this outline and shoot your video immediately.")
    add_tip(doc, "Ideas get smarter over time as the Insights engine learns what works for you.")

    # Video 8
    add_video(doc, "8", "AI Generation", "5 min",
        "Complete guide to the Generate tab: AI thumbnails with image upload, AI images, AI B-roll video.")
    add_visual(doc, "Navigate to Generate tab. Show the three sub-tabs.")
    add_narration(doc, "The Generate tab is where AI creates visual assets. There are three modes: AI Thumbnails, AI Images, and AI B-roll Video. Let's start with thumbnails.")
    add_visual(doc, "With AI Thumbnail selected, enter a video title. Show the Upload your own image section.")
    add_narration(doc, "For thumbnails, enter your video title. Now here is the exciting part, you can upload your own photo and AI transforms it into a stylized thumbnail. Click to upload and select an image, like a photo of your dog.")
    add_visual(doc, "Upload a photo. Show the preview and style strength slider.")
    add_narration(doc, "Once uploaded, you see a preview and a style strength slider. At 10 to 30 percent, your photo stays very recognizable with minimal AI styling. At 30 to 50 percent, you get a balanced mix. Above 50 percent, the AI transforms the image more dramatically. The default 35 percent is a great starting point.")
    add_visual(doc, "Click Generate from your photo. Show result in the grid.")
    add_narration(doc, "Click Generate from your photo and the system uses Replicate SDXL image-to-image model to transform your upload. This takes about 30 to 60 seconds. Without an upload, AI generates from scratch using the built-in image generator, no API key needed.")
    add_visual(doc, "Switch to AI B-roll Video tab. Show the Replicate warning.")
    add_narration(doc, "AI B-roll Video generates actual video clips from text prompts. This requires a Replicate API token which you add in the API Keys tab. Generation takes 2 to 5 minutes since it creates actual video frames.")
    add_tip(doc, "Show the model name and status badge on each generated asset.")

    # Video 9 - NEW: Script Analyzer
    add_video(doc, "9", "Script Analyzer - Reverse-Engineer Viral Videos", "5 min",
        "Paste any viral video URL and AI breaks down the script, psychology triggers, framework, and viral score. Adapt it for your niche in one click.")
    add_visual(doc, "Navigate to the Analyzer tab. Show the URL input and mode selector.")
    add_narration(doc, "The Script Analyzer is one of the most powerful features in ContentForge. Paste any viral video URL and AI reverse-engineers exactly why it works. It breaks down the script into its component parts, identifies the psychology triggers, extracts a reusable framework, and scores its viral potential.")
    add_visual(doc, "Show the Single Video / Bulk Analysis toggle.")
    add_narration(doc, "You can analyze a single video or up to 20 at once in bulk mode. Let's start with single video analysis.")
    add_visual(doc, "Type a YouTube URL in the input field. Check the auto-adapt checkbox. Click Analyze Video.")
    add_narration(doc, "Paste a YouTube URL, check Auto-adapt script for my niche, and click Analyze Video. The system fetches the transcript, sends it to AI for analysis, and generates a complete breakdown. This takes about 30 to 60 seconds.")
    add_visual(doc, "Show the analysis result appearing with four color-coded sections: Hook, Pattern Interrupt, Body, CTA.")
    add_narration(doc, "Each analysis shows four color-coded sections. The Hook is the exact opening text from the first 3 seconds. The Pattern Interrupt is what breaks the viewer's scroll expectation. The Body is the main content. And the CTA is the call to action.")
    add_visual(doc, "Scroll down to show psychology triggers as colored badges.")
    add_narration(doc, "Below the script breakdown, you see psychology trigger badges. These are the psychological techniques the video uses: urgency, scarcity, social proof, curiosity, FOMO, authority, reciprocity, and loss aversion. Each trigger is color-coded so you can spot patterns across multiple videos.")
    add_visual(doc, "Show the framework section with numbered steps.")
    add_narration(doc, "Next is the Framework. This is the reusable structure. For example, Problem, Agitation, Solution, Proof, CTA. You can save this framework to your Framework Library and use it as a template for your own content.")
    add_visual(doc, "Show the viral score and analysis notes.")
    add_narration(doc, "The viral score rates the video from 0 to 100. The analysis notes explain why the video works, what makes it effective, and what you can learn from it.")
    add_visual(doc, "Show the adapted script in the orange box.")
    add_narration(doc, "If you checked Auto-adapt, you also get an adapted script in this orange box. AI takes the original framework and psychology, but swaps the content to fit your niche. Same proven structure, your topic. You can film this script immediately.")
    add_visual(doc, "Click Adapt for My Niche button. Click Save as Framework button.")
    add_narration(doc, "Two action buttons appear on each analysis. Adapt for My Niche generates a new adapted version if you didn't auto-adapt. Save as Framework adds the reusable template to your Framework Library for future use.")
    add_visual(doc, "Switch to Bulk Analysis mode. Show pasting multiple URLs.")
    add_narration(doc, "Bulk Analysis mode lets you paste up to 20 URLs at once, one per line. The system analyzes each one and sorts results by viral score. This is perfect for finding patterns across multiple viral videos in your niche.")
    add_tip(doc, "YouTube videos work best for transcript extraction. TikTok and Instagram are supported but less reliable.")

    # Video 10 - NEW: Framework Library
    add_video(doc, "10", "Framework Library - Reusable Templates", "3 min",
        "How to use saved frameworks as starting points for your content.")
    add_visual(doc, "Navigate to the Frameworks tab. Show the grid of saved frameworks.")
    add_narration(doc, "The Framework Library is your collection of reusable script templates extracted from viral videos. Every time you click Save as Framework in the Analyzer, it appears here.")
    add_visual(doc, "Point to each element on a framework card: name, description, steps, psychology triggers.")
    add_narration(doc, "Each framework card shows the name, a description of why it works, the numbered steps that make up the structure, and the psychology triggers it uses. For example, a framework might be Problem, Agitation, Solution, Proof, CTA with urgency and social proof triggers.")
    add_visual(doc, "Expand the example script section.")
    add_narration(doc, "Click View example script to see a filled-in version. This shows you exactly how to apply the framework to real content. Use this as a starting template and swap in your own product or topic.")
    add_visual(doc, "Point to the use count and source URL.")
    add_narration(doc, "The use count tracks how many times you have adapted this framework. The source URL links back to the original viral video you extracted it from, so you can revisit the inspiration anytime.")
    add_visual(doc, "Click Delete framework on one card.")
    add_narration(doc, "Delete a framework anytime if it's no longer useful. Your framework library grows over time, giving you a personal collection of proven viral structures specific to your niche.")
    add_tip(doc, "Aim to collect 5-10 frameworks. That gives you enough variety to test different structures and find what works best for your audience.")

    # Video 11
    add_video(doc, "11", "Social Accounts & API Keys", "6 min",
        "How to add API credentials and connect social media accounts via OAuth.")
    add_visual(doc, "Navigate to API Keys tab. Show the 5 platform groups.")
    add_narration(doc, "Before you can publish to social platforms, you need to add API credentials. Go to the API Keys tab. You see five platform groups: YouTube, TikTok, Instagram plus Facebook, X, and Replicate for AI video generation.")
    add_visual(doc, "Click Get keys next to YouTube. Explain the process.")
    add_narration(doc, "Click Get keys to open the platform's developer portal. For YouTube, create a Google Cloud Project, enable the YouTube Data API, and create OAuth credentials. The system shows the exact redirect URL to use.")
    add_visual(doc, "Paste credentials and click Save. Show the Configured badge.")
    add_narration(doc, "Paste your Client ID and Secret and click Save. Values are encrypted before storage. You see a masked preview. Once saved, the card shows a green Configured badge.")
    add_visual(doc, "Navigate to Social tab. Show YouTube now showing API Ready with Connect button.")
    add_narration(doc, "Now go to the Social tab. YouTube shows API Ready with a Connect button. Click it to start the OAuth flow. You are redirected to Google's login page. Authorize the app and you are sent back.")
    add_visual(doc, "Show connected account with handle. Point to Needs Keys badges on other platforms.")
    add_narration(doc, "Your YouTube account now shows as connected. Repeat for each platform. TikTok, Instagram, Facebook, and X each have their own developer portals with links provided.")
    add_tip(doc, "Instagram and Facebook require videos hosted on a public URL for publishing. You will need S3 or Cloudinary storage.")

    # Video 12
    add_video(doc, "12", "Scheduling & Calendar", "4 min",
        "How to schedule posts at optimal times and view them in the content calendar.")
    add_visual(doc, "From Library, click Publish on a ready video. Show the When to publish section.")
    add_narration(doc, "When you publish a video, you have three scheduling options. Click Publish on any ready video and scroll to the When to publish section.")
    add_visual(doc, "Point to the three buttons: Publish now, Optimal times, Pick time.")
    add_narration(doc, "Publish immediately, schedule at AI-determined optimal times, or pick a specific time. Optimal times uses industry-standard peak engagement hours for each platform. YouTube at 3 PM Eastern, TikTok at 9 AM, noon, or 7 PM, Instagram at 11 AM, 2 PM, or 7 PM.")
    add_visual(doc, "Select Optimal times and click Schedule optimal. Show success message.")
    add_narration(doc, "Once scheduled, the post appears in the Scheduled tab with a blue badge. Vercel Cron checks every minute and when the time comes, it automatically publishes.")
    add_visual(doc, "Navigate to Calendar tab. Show the month grid.")
    add_narration(doc, "The Calendar tab gives you a visual month view of all scheduled and published posts. Blue dots for scheduled, green for published, red for failed. Click month navigation to look ahead or review past performance.")
    add_tip(doc, "The agent can also schedule posts. Just ask it to schedule my ready videos at optimal times.")

    # Video 13
    add_video(doc, "13", "Trend Research", "3 min",
        "How to use the Trends tab to discover what's currently trending.")
    add_visual(doc, "Navigate to Trends tab. Show the Refresh trends button.")
    add_narration(doc, "The Trends tab shows what's currently trending in your niche across TikTok, Instagram, YouTube, and X. Trends are refreshed daily automatically, but you can trigger a refresh manually.")
    add_visual(doc, "Click Refresh trends. Show results appearing.")
    add_narration(doc, "Click Refresh trends and the system performs web searches for trending content in your niche, then uses AI to extract and score the trends. This takes about 30 to 60 seconds.")
    add_visual(doc, "Show the trends list with type icons, content, summary, score.")
    add_narration(doc, "Each trend shows the platform, a type icon (hashtag, sound, format, or topic), the trend content, a summary explaining why it's trending, a relevance score from 0 to 100, and when it was discovered.")
    add_visual(doc, "Click a platform filter button.")
    add_narration(doc, "Filter by platform using these buttons. Trend data also feeds into the ideation engine, so generated ideas are grounded in what's actually trending right now.")
    add_tip(doc, "Trends feed automatically into video analysis. When a video is processed, AI includes current trending hashtags in its suggestions.")

    # Video 14
    add_video(doc, "14", "Analytics Dashboard", "4 min",
        "How to view cross-platform analytics and understand performance.")
    add_visual(doc, "Navigate to Analytics tab. Show the four stat cards.")
    add_narration(doc, "The Analytics tab pulls performance metrics from all your connected social platforms into one unified view. At the top, four stat cards: Total Views, Total Likes, Comments, and Shares.")
    add_visual(doc, "Scroll to Per-Platform Breakdown.")
    add_narration(doc, "Below the totals, a per-platform breakdown showing post count and views per platform. This helps identify which platforms perform best for your content.")
    add_visual(doc, "Scroll to Top Performing Videos.")
    add_narration(doc, "The Top Performing Videos section ranks your content by views. This is gold for understanding what resonates with your audience. Double down on what works.")
    add_visual(doc, "Click Refresh metrics.")
    add_narration(doc, "Metrics refresh automatically every hour, but you can click Refresh metrics to pull fresh data immediately.")
    add_tip(doc, "Analytics data feeds into the Insights engine which learns patterns and feeds them back into ideation.")

    # Video 15
    add_video(doc, "15", "Insights Engine", "3 min",
        "How the learning loop works. AI analyzes your analytics to generate actionable insights.")
    add_visual(doc, "Navigate to Insights tab. Show the Refresh insights button.")
    add_narration(doc, "The Insights tab is where the system learns from your performance. AI analyzes your last 30 videos and analytics data to identify patterns, opportunities, underperformers, and strategic recommendations.")
    add_visual(doc, "Click Refresh insights. Show results.")
    add_narration(doc, "Click Refresh insights and AI generates 5 to 8 actionable insights. Each is color-coded: green Patterns show what's working, blue Opportunities highlight untapped potential, red Underperformers flag what's failing, amber Recommendations give strategic advice.")
    add_visual(doc, "Expand the data section on one insight.")
    add_narration(doc, "Each insight includes supporting data, actual numbers from your analytics. The confidence score shows how certain the AI is about the pattern. Insights feed directly into the ideation engine, creating a true learning loop.")
    add_tip(doc, "Insights refresh daily via cron job, staying current as you publish more content.")

    # Video 16
    add_video(doc, "16", "Auto-Comment Replies", "5 min",
        "How to fetch comments, review AI-suggested replies, and approve or edit before posting.")
    add_visual(doc, "Navigate to Comments tab. Show the Fetch new comments button.")
    add_narration(doc, "The Comments tab is your comment management center. AI fetches comments from your published posts and generates replies in your voice.")
    add_visual(doc, "Click Fetch new comments. Show comments appearing.")
    add_narration(doc, "Click Fetch new comments and the system pulls recent comments from YouTube, Instagram, Facebook, and X. For each new comment, AI generates a suggested reply based on your Voice Profile.")
    add_visual(doc, "Show a comment card with AI suggested reply in blue box.")
    add_narration(doc, "Each comment card shows the author, platform, comment text, and the video it was posted on. Below that, in the blue box, is the AI-suggested reply written in your voice, matching your persona and tone.")
    add_visual(doc, "Point to the three buttons: Approve and Post, Edit, Ignore.")
    add_narration(doc, "You have three options. Approve and Post sends the AI reply immediately. Edit lets you modify before posting. Ignore marks the comment as skipped.")
    add_visual(doc, "Show filter buttons: Pending, Replied, Ignored, All.")
    add_narration(doc, "Use the filter tabs to view comments by status. If you enable auto-reply mode in the Voice tab, AI posts replies automatically without approval.")
    add_tip(doc, "Comments are checked 4 times daily by a cron job, so you never miss new engagement.")

    # Video 17
    add_video(doc, "17", "Voice Profile Setup", "3 min",
        "How to configure your AI voice profile for consistent replies and scripts.")
    add_visual(doc, "Navigate to Voice tab. Show the voice profile form.")
    add_narration(doc, "The Voice tab is where you define your AI persona, the voice that generates comment replies, voiceover scripts, and other AI-written content. Getting this right makes everything sound like YOU.")
    add_visual(doc, "Point to the Persona field.")
    add_narration(doc, "Start with the Persona field. Describe yourself in a sentence, like friendly dog mom who loves puns or energetic fitness coach who motivates beginners.")
    add_visual(doc, "Point to Tone, Signature Phrases, and Phrases to Avoid.")
    add_narration(doc, "Set your Tone, like warm, witty, enthusiastic. Add Signature Phrases you use frequently, like pawsome, fur-real, zoomies. And add Phrases to Avoid, words you never want AI to use.")
    add_visual(doc, "Point to Reply Mode and Reply Length dropdowns.")
    add_narration(doc, "Reply Mode has two options. Suggest generates replies for manual approval, recommended. Auto posts replies immediately. Reply Length controls how long AI replies are: Short, Medium, or Long.")
    add_visual(doc, "Click Save Voice Profile.")
    add_narration(doc, "Click Save and your voice profile immediately affects all AI-generated content. Update it anytime as your brand evolves.")
    add_tip(doc, "Experiment with different personas to see which generates the best replies.")

    # Video 18
    add_video(doc, "18", "Competitor Monitoring", "4 min",
        "How to add competitors, get viral alerts, and use AI suggestions.")
    add_visual(doc, "Navigate to Competitors tab. Show the add competitor form.")
    add_narration(doc, "The Competitors tab lets you track other creators in your niche. When their posts go viral, you get an alert with AI suggestions on how to create your own version.")
    add_visual(doc, "Select youtube, type a handle, click Add.")
    add_narration(doc, "Select the platform and type the competitor's handle without the @ symbol. Click Add and they are tracked. The system checks their posts daily.")
    add_visual(doc, "Click Check now. Show results.")
    add_narration(doc, "Click Check now to fetch their latest posts immediately. AI analyzes each post, estimates a viral score, and flags posts scoring 70 or above as viral alerts.")
    add_visual(doc, "Show the Viral alerts section with AI suggestions.")
    add_narration(doc, "Viral alerts appear in this orange section. Each alert shows the post title, competitor, viral score, and an AI suggestion, like Create a similar video about X but add Y twist. This turns competitor research into actionable content ideas.")
    add_tip(doc, "Add 3-5 competitors for good coverage without alert fatigue.")

    # Video 19
    add_video(doc, "19", "Brand Kit", "3 min",
        "How to set up your brand logo, colors, and fonts for consistent video branding.")
    add_visual(doc, "Navigate to Brand Kit tab. Show the logo upload section.")
    add_narration(doc, "The Brand Kit tab stores your visual identity, logo, colors, and fonts, so every video gets consistent branding automatically.")
    add_visual(doc, "Click Upload logo. Show preview appearing.")
    add_narration(doc, "Start by uploading your logo. PNG with transparency works best. Click Upload logo and select your file.")
    add_visual(doc, "Show the color pickers and font dropdown.")
    add_narration(doc, "Set your brand colors: Primary, Secondary, and Accent. Choose your Font Family from the list. These affect burned-in captions and text overlays.")
    add_visual(doc, "Show Watermark Position and Size controls. Click Save Brand Kit.")
    add_narration(doc, "Control where your logo appears and how large it is. Click Save Brand Kit and every future video automatically uses your logo as a watermark.")
    add_tip(doc, "The agent can update brand kit settings. Just ask it to set my primary color to blue.")

    # Video 20
    add_video(doc, "20", "Content Repurposing", "4 min",
        "How to turn one long video into multiple short clips automatically.")
    add_visual(doc, "Navigate to Library. Find a video longer than 90 seconds. Point to the Repurpose button.")
    add_narration(doc, "Content repurposing is one of the most powerful features. If you have a long video, you can turn it into 5 short clips for TikTok, Reels, and Shorts with one click.")
    add_visual(doc, "Click Repurpose. Show confirmation.")
    add_narration(doc, "Find any video longer than 90 seconds and click the purple Repurpose button. AI analyzes the transcript, identifies the 5 most engaging moments, and extracts each as a separate clip.")
    add_visual(doc, "Wait, then refresh Library. Show new clip cards with purple Clip badges.")
    add_narration(doc, "After a few minutes, refresh your Library. New video cards appear with a purple Clip badge. Each clip has its own title, viral score, and timestamp showing where it was cut from. Each clip is fully processed with captions, formats, and hashtags.")
    add_visual(doc, "Show the agent. Ask it to repurpose a video.")
    add_narration(doc, "You can also trigger repurposing through the AI Agent. Just ask it to repurpose my latest video into 5 clips.")
    add_tip(doc, "Repurposing works best on videos with clear speech since AI needs the transcript.")

    # Video 21
    add_video(doc, "21", "Subtitle Translation", "3 min",
        "How to translate your video captions into 12 languages.")
    add_visual(doc, "In Library, click Translate on a ready video. Show the dialog.")
    add_narration(doc, "If you want to reach international audiences, ContentForge can translate your captions into 12 languages with one click.")
    add_visual(doc, "Show the 12 language options. Select 3. Click Translate.")
    add_narration(doc, "The Translate dialog shows 12 supported languages: Spanish, French, German, Italian, Portuguese, Dutch, Russian, Japanese, Korean, Chinese, Arabic, and Hindi. Select multiple at once and click Translate.")
    add_visual(doc, "Show translations appearing in Existing translations section.")
    add_narration(doc, "Translated text appears in the Existing translations section. Each translation preserves your tone, emojis, and brand voice. The agent can also translate. Just ask it to translate my latest video to Spanish and French.")
    add_tip(doc, "Use translations when publishing to YouTube, which supports multiple language tracks.")

    # Video 22
    add_video(doc, "22", "Hook A/B Testing", "3 min",
        "How to generate hook variants and learn which opening styles work best.")
    add_visual(doc, "Explain the concept of hooks.")
    add_narration(doc, "The first 3 seconds of a video decide whether someone scrolls past or watches to the end. That opening is called a hook. ContentForge generates 5 different hook variants for any video, each using a different psychological technique.")
    add_visual(doc, "Ask the agent: 'Generate 5 hooks for my latest video'")
    add_narration(doc, "To generate hooks, ask the AI Agent to generate 5 hooks for your latest video. The agent analyzes the transcript and creates 5 variants: Question, Bold Claim, Curiosity Gap, Stat, and Story.")
    add_visual(doc, "Show the hooks with predicted scores.")
    add_narration(doc, "Each hook includes a predicted viral score and reasoning. Use the highest-scoring hook as the text overlay in your video's opening 3 seconds.")
    add_visual(doc, "Ask the agent: 'What hook styles work best for me?'")
    add_narration(doc, "After your posts have been live for 24 hours, the system checks which hook style performed best. Ask the agent what hook styles work best for me and it shows performance breakdowns by style. This data-driven approach takes the guesswork out of your video openings.")
    add_tip(doc, "Hooks are stored per video, so you can compare performance across multiple videos over time.")

    # Video 23
    add_video(doc, "23", "Multi-Format Publishing", "3 min",
        "How the system auto-generates 3 formats and publishes with platform-specific captions.")
    add_visual(doc, "In Library, show the format badges: 16:9, 9:16, 1:1.")
    add_narration(doc, "Every video you upload is automatically generated in three formats: 16:9 horizontal for YouTube and Facebook, 9:16 vertical for TikTok and Reels, and 1:1 square for Instagram feed. You do not need to do anything, it is all automatic.")
    add_visual(doc, "Click Publish. Show platform selection.")
    add_narration(doc, "When you publish, you select which platforms to post to. The system automatically uses the correct format for each platform. But it is not just the video format that is optimized.")
    add_visual(doc, "Explain per-platform caption generation.")
    add_narration(doc, "When you publish, AI generates platform-specific captions. YouTube gets SEO-optimized title and description with 10 to 15 hashtags. TikTok gets a punchy title with emojis and 4 to 6 trending hashtags. Instagram gets an emoji-rich caption with 15 to 30 hashtags. X gets a single tweet under 280 characters. Each platform gets content optimized for its culture.")
    add_visual(doc, "End with a summary shot of the full workflow.")
    add_narration(doc, "And that is the ContentForge workflow. Upload once, AI handles everything else. Editing, captions, voiceover, multi-format generation, platform-specific copy, optimal scheduling, and publishing across all your social platforms. One video becomes optimized content for five platforms, on autopilot.")
    add_tip(doc, "End with a call to action: Start automating your content at ContentForge dot com.")

    # ============ PART 2: FIELD GUIDE ============
    doc.add_page_break()
    add_heading(doc, "PART 2: FIELD-BY-FIELD GUIDE", level=1, color=RGBColor(0xE8, 0x6A, 0x00))
    add_para(doc, "Every input, dropdown, toggle, and button explained. Use as your reference manual.", italic=True, size=11)

    # Dashboard
    doc.add_page_break()
    add_heading(doc, "Dashboard Tab", level=2)
    add_field(doc, "Total Videos", "Stat card", "Dashboard", "Total count of all videos in your library regardless of status.")
    add_field(doc, "Processing", "Stat card", "Dashboard", "Videos currently being edited or analyzed.")
    add_field(doc, "Ready", "Stat card", "Dashboard", "Videos that completed processing and are ready to publish.")
    add_field(doc, "Scheduled", "Stat card", "Dashboard", "Posts scheduled but not yet published.")
    add_field(doc, "Published", "Stat card", "Dashboard", "Videos successfully published to at least one platform.")
    add_field(doc, "Avg Viral Score", "Stat card", "Dashboard", "Average viral score (0-100) across all scored videos.")
    add_field(doc, "Quick Actions", "Buttons", "Dashboard", "Shortcuts to common tasks. Click to navigate.")
    add_field(doc, "Connected Accounts", "Display", "Dashboard", "Shows how many social accounts are connected and on which platforms.")

    # Upload
    doc.add_page_break()
    add_heading(doc, "Upload Tab", level=2)
    add_field(doc, "Video Drop Zone", "File input", "Upload tab", "Click or drag video files. Accepts MP4, MOV, WebM. Multiple files allowed.", "Drag 3 MP4 files here to upload all 3.")
    add_field(doc, "Upload Photos Instead", "Button", "Upload tab", "Opens file picker for images. Photos are converted to video with Ken Burns effect.", "Click to select 4 photos for a slideshow.")
    add_field(doc, "Seconds per Photo", "Slider (2-10s)", "Upload tab (photos)", "How long each photo displays. Default: 5 seconds.", "3s for fast-paced, 7s for dramatic.")
    add_field(doc, "Transition", "Slider (0.3-2s)", "Upload tab (photos)", "Crossfade duration between photos. Default: 0.7 seconds.")
    add_field(doc, "Custom Voiceover Script", "Textarea", "Upload tab (photos)", "Optional script for AI voiceover. If blank, AI generates one.", "Type 'Meet Max, the goodest boy!' to use that exact text.")
    add_field(doc, "Burn AI Captions", "Toggle", "Upload tab", "When ON, AI transcribes audio and burns subtitles onto the video.", "ON = subtitles appear on the final video.", "Keep ON for most content.")
    add_field(doc, "Auto-trim Silence", "Toggle", "Upload tab", "When ON, removes silent segments from the video.", "ON = 60s video with 10s silence becomes 50s.", "Good for talking-head videos.")
    add_field(doc, "Watermark", "Dropdown", "Upload tab", "Select a PNG logo to overlay. Upload logos in Assets tab first.", "Select MyLogo.png to add it to every frame.")
    add_field(doc, "Watermark Position", "Dropdown", "Upload tab", "Where the watermark appears: Top Left, Top Right, Bottom Left, Bottom Right, Center.", "Bottom Right is standard.")
    add_field(doc, "Background Music", "Dropdown", "Upload tab", "Select music to mix under your video. Upload in Assets tab first.", "Select upbeat_track.mp3 for background music.", "Royalty-free only.")
    add_field(doc, "Intro Clip", "Dropdown", "Upload tab", "Video clip to prepend. Upload in Assets tab first.", "Select intro_3sec.mp4 for a branded intro.", "Keep intros under 5 seconds.")
    add_field(doc, "Outro Clip", "Dropdown", "Upload tab", "Video clip to append. Upload in Assets tab first.")
    add_field(doc, "Enable AI Voiceover", "Toggle", "Upload tab", "When ON, AI generates a script and synthesizes it as audio.", "Auto-enabled for photo uploads.")
    add_field(doc, "Voice", "Dropdown", "Upload tab", "Select from 7 AI voices with different genders and tones.", "Tongtong = warm neutral. Tianmei = cheerful female.", "Test different voices to find your match.")
    add_field(doc, "Tone", "Text input", "Upload tab", "Describe the tone for the AI script.", "funny, energetic, engaging", "Be specific. 'Funny with dog puns' is better than just 'funny'.")
    add_field(doc, "Replace Original Audio", "Toggle", "Upload tab", "ON = voiceover replaces all audio. OFF = voiceover mixes on top.", "For photos, always ON. For videos, OFF usually sounds better.")

    # Library
    doc.add_page_break()
    add_heading(doc, "Library Tab", level=2)
    add_field(doc, "Viral Score Circle", "Display", "Library card top right", "AI score 0-100. Green=80+, Amber=50-79, Red=below 50.")
    add_field(doc, "Status Badge", "Display", "Library card top left", "Pipeline status: pending, editing, ready, published, failed.")
    add_field(doc, "Clip Badge", "Display", "Library card top left", "Purple badge for clips extracted via Repurpose.")
    add_field(doc, "Format Badges", "Display", "Library card", "Green badges: 16:9, 9:16, 1:1 showing generated formats.")
    add_field(doc, "Publish", "Button", "Library card", "Opens publish dialog. Only appears when video is ready.")
    add_field(doc, "Repurpose", "Button", "Library card", "Purple button. Cuts long videos into 5 clips. Only on videos over 90 seconds.")
    add_field(doc, "Translate", "Button", "Library card", "Opens translate dialog for 12 languages.")
    add_field(doc, "Retry", "Button", "Library card", "Reprocesses a failed video. Only appears when status is failed.")
    add_field(doc, "Delete", "Button", "Library card", "Permanently deletes the video and all files.")

    # Ideas
    doc.add_page_break()
    add_heading(doc, "Ideas Tab", level=2)
    add_field(doc, "Generate 5 ideas", "Button", "Ideas tab top right", "Triggers AI to generate 5 content ideas from trends + analytics. Takes 15-20s.", "Click to get 5 tailored ideas.", "Generates different ideas each time.")
    add_field(doc, "Idea Card - Title", "Display", "Ideas tab", "AI-generated scroll-stopping title for a content idea.")
    add_field(doc, "Idea Card - Predicted Viral Score", "Display", "Ideas tab", "AI prediction 0-100. Color-coded.")
    add_field(doc, "Idea Card - Format", "Display", "Ideas tab", "Recommended aspect ratio: 9:16, 16:9, or 1:1.")
    add_field(doc, "Idea Card - Script Outline", "Display", "Ideas tab", "Hook + Body + CTA breakdown. Production-ready.")
    add_field(doc, "Dismiss", "Button", "Ideas tab", "Removes an idea from the list.")

    # Generate
    doc.add_page_break()
    add_heading(doc, "Generate Tab", level=2)
    add_field(doc, "AI Thumbnail / AI Image / AI B-roll", "Tab selector", "Generate tab", "Switch between 3 generation modes.")
    add_field(doc, "Video title", "Text input", "Generate tab (thumbnail)", "Title to base the thumbnail on.")
    add_field(doc, "Upload your own image", "File input", "Generate tab (thumbnail)", "Optional. Upload a photo and AI transforms it via Replicate SDXL img2img.", "Upload my_dog.jpg for a custom thumbnail.", "Requires Replicate API token.")
    add_field(doc, "Style Strength", "Slider (10-80%)", "Generate tab (thumbnail)", "How much AI changes your photo. Lower=more recognizable. Higher=more transformation.", "35% default = balanced.", "Start at 35%.")
    add_field(doc, "Prompt", "Textarea", "Generate tab (broll)", "Describe the video clip you want AI to generate.", "cinematic shot of a golden retriever running on a beach", "Requires Replicate API token. Takes 2-5 minutes.")

    # Script Analyzer - NEW
    doc.add_page_break()
    add_heading(doc, "Script Analyzer Tab", level=2)
    add_field(doc, "Single Video / Bulk Analysis", "Toggle", "Analyzer tab", "Switch between analyzing one URL or up to 20 at once.", "Select Bulk to paste 10 URLs and sort by viral score.")
    add_field(doc, "Video URL", "Text input", "Analyzer tab (single)", "Paste a YouTube, TikTok, or Instagram video URL to analyze.", "https://www.youtube.com/watch?v=abc123", "YouTube works best for transcript extraction.")
    add_field(doc, "Video URLs", "Textarea", "Analyzer tab (bulk)", "Paste multiple URLs, one per line. Max 20.", "One URL per line for bulk analysis.")
    add_field(doc, "Auto-adapt script for my niche", "Checkbox", "Analyzer tab", "When checked, AI also generates an adapted version of the script for your niche.", "Checked = you get both the original breakdown AND an adapted version.", "Save time by checking this. You can always adapt later if you forget.")
    add_field(doc, "Analyze Video / Analyze N Videos", "Button", "Analyzer tab", "Triggers the analysis. Single takes 30-60s. Bulk takes 2-5 min.")
    add_field(doc, "Hook", "Display (amber box)", "Analyzer tab result", "The exact opening text from the first 0-3 seconds of the video.", "'You won't believe what this dog did next!'", "The hook is the most important part. If the first 3 seconds don't grab attention, the rest doesn't matter.")
    add_field(doc, "Pattern Interrupt", "Display (purple box)", "Analyzer tab result", "What breaks the viewer's scroll expectation. The element of surprise.", "'Suddenly cuts to the dog wearing sunglasses'", "Pattern interrupts are what make viewers stop scrolling and watch.")
    add_field(doc, "Body", "Display (blue box)", "Analyzer tab result", "The main content or message of the video.")
    add_field(doc, "CTA", "Display (green box)", "Analyzer tab result", "The call to action. What the video asks viewers to do.", "'Follow for more daily dog content!'", "Strong CTAs drive follows and engagement.")
    add_field(doc, "Psychology Triggers", "Display (colored badges)", "Analyzer tab result", "Psychological techniques used: urgency, scarcity, social proof, curiosity, FOMO, authority, reciprocity, loss aversion.", "Badges: 'urgency', 'social_proof', 'curiosity'", "Color-coded for easy pattern spotting across multiple videos.")
    add_field(doc, "Framework", "Display (numbered steps)", "Analyzer tab result", "The reusable structure of the video.", "Steps: Hook, Agitation, Solution, Proof, CTA", "Save frameworks you like to the Framework Library for reuse.")
    add_field(doc, "Viral Score", "Display", "Analyzer tab result", "AI's estimate of the video's viral potential (0-100).")
    add_field(doc, "Analysis Notes", "Display", "Analyzer tab result", "AI's explanation of why the video works.")
    add_field(doc, "Adapted Script", "Display (orange box)", "Analyzer tab result", "AI-adapted version of the script for your niche. Same psychology, different content.", "Shows: HOOK: [adapted text], BODY: [adapted text], CTA: [adapted text]", "Film this script immediately. It's ready to use.")
    add_field(doc, "Adapt for My Niche", "Button", "Analyzer tab result", "Generates a new adapted version if you didn't auto-adapt.")
    add_field(doc, "Save as Framework", "Button", "Analyzer tab result", "Saves the reusable framework to your Framework Library.")

    # Framework Library - NEW
    doc.add_page_break()
    add_heading(doc, "Framework Library Tab", level=2)
    add_field(doc, "Framework Card", "Display", "Frameworks tab", "Shows framework name, description, steps, psychology triggers, example script, source URL, use count.")
    add_field(doc, "Steps", "Display (numbered)", "Frameworks tab card", "The sequential structure of the framework.", "1. Hook, 2. Agitation, 3. Solution, 4. Proof, 5. CTA")
    add_field(doc, "Psychology Triggers", "Display (badges)", "Frameworks tab card", "Which psychological techniques this framework uses.")
    add_field(doc, "View example script", "Expandable", "Frameworks tab card", "Click to see a filled-in example of the framework in action.")
    add_field(doc, "Source URL", "Link", "Frameworks tab card", "Links to the original viral video this framework was extracted from.")
    add_field(doc, "Use Count", "Display", "Frameworks tab card", "How many times you have adapted this framework.", "Shows 'Used 5x' if you have adapted it 5 times.", "Popular frameworks rise to the top automatically.")
    add_field(doc, "Delete framework", "Button", "Frameworks tab card", "Removes the framework from your library.")

    # Social
    doc.add_page_break()
    add_heading(doc, "Social Tab", level=2)
    add_field(doc, "Platform Status Badge", "Display", "Social tab", "Shows 'API Ready' (green) or 'Needs Keys' (amber).")
    add_field(doc, "Connect", "Button", "Social tab", "Starts OAuth flow. Redirects to platform login. Only appears when API is ready.")
    add_field(doc, "Disconnect", "Button", "Social tab", "Marks account as disconnected.")

    # Scheduled
    doc.add_page_break()
    add_heading(doc, "Scheduled Tab", level=2)
    add_field(doc, "Summary Cards", "Display", "Scheduled tab", "Four cards: Scheduled, In progress, Published (24h), Failed (24h).")
    add_field(doc, "Cancel", "Button", "Scheduled tab", "Cancels a scheduled post before it publishes.")

    # Calendar
    doc.add_page_break()
    add_heading(doc, "Calendar Tab", level=2)
    add_field(doc, "Month Navigation", "Buttons", "Calendar tab", "Navigate between months.")
    add_field(doc, "Day Cells", "Display grid", "Calendar tab", "Calendar grid. Today highlighted in orange. Posts shown as colored dots.", "Blue=scheduled, Green=published, Red=failed, Amber=uploading")

    # Trends
    doc.add_page_break()
    add_heading(doc, "Trends Tab", level=2)
    add_field(doc, "Refresh trends", "Button", "Trends tab", "Triggers web search for trending content. Takes 30-60s.", "Also auto-refreshes daily via cron.")
    add_field(doc, "Platform Filter", "Buttons", "Trends tab", "Filter by platform: All, TikTok, Instagram, YouTube, X.")
    add_field(doc, "Score", "Display", "Trends tab card", "Relevance score 0-100. Higher = more relevant to your niche.")

    # Analytics
    doc.add_page_break()
    add_heading(doc, "Analytics Tab", level=2)
    add_field(doc, "Total Views/Likes/Comments/Shares", "Stat cards", "Analytics tab", "Aggregate metrics across all platforms.")
    add_field(doc, "Refresh metrics", "Button", "Analytics tab", "Pulls fresh metrics. Also auto-refreshes hourly.")
    add_field(doc, "Per-Platform Breakdown", "Display", "Analytics tab", "Post count and views per platform.")
    add_field(doc, "Top Performing Videos", "Display", "Analytics tab", "Top 10 videos ranked by views.")

    # Insights
    doc.add_page_break()
    add_heading(doc, "Insights Tab", level=2)
    add_field(doc, "Refresh insights", "Button", "Insights tab", "AI analyzes last 30 videos + analytics. Takes ~20s.")
    add_field(doc, "Insight Type", "Display (color-coded)", "Insights tab", "Pattern (green), Opportunity (blue), Underperformer (red), Recommendation (amber).")
    add_field(doc, "Confidence", "Display", "Insights tab", "How certain AI is (0-100%).")

    # Comments
    doc.add_page_break()
    add_heading(doc, "Comments Tab", level=2)
    add_field(doc, "Fetch new comments", "Button", "Comments tab", "Pulls comments from published posts and generates AI replies.", "Also auto-runs 4x daily via cron.")
    add_field(doc, "Filter (Pending/Replied/Ignored/All)", "Buttons", "Comments tab", "Filter comments by reply status.")
    add_field(doc, "Approve & Post", "Button", "Comments tab card", "Posts the AI-suggested reply immediately.")
    add_field(doc, "Edit", "Button", "Comments tab card", "Opens the reply for editing before posting.")
    add_field(doc, "Ignore", "Button", "Comments tab card", "Marks comment as ignored. No reply posted.")

    # Competitors
    doc.add_page_break()
    add_heading(doc, "Competitors Tab", level=2)
    add_field(doc, "Platform Dropdown", "Dropdown", "Competitors tab", "Select platform: YouTube, Instagram, TikTok, X.")
    add_field(doc, "Handle Input", "Text input", "Competitors tab", "Type competitor handle without @.", "Type 'mrbeast' not '@mrbeast'.")
    add_field(doc, "Add", "Button", "Competitors tab", "Adds competitor to tracked list.")
    add_field(doc, "Check now", "Button", "Competitors tab", "Fetches latest posts from all competitors.", "Also auto-runs daily at 7am via cron.")
    add_field(doc, "Remove", "Button", "Competitors tab", "Stops tracking a competitor.")
    add_field(doc, "Viral Alerts", "Display", "Competitors tab", "Competitor posts scoring 70+ with AI suggestions for your own version.")

    # Brand Kit
    doc.add_page_break()
    add_heading(doc, "Brand Kit Tab", level=2)
    add_field(doc, "Upload logo", "Button", "Brand Kit tab", "Upload PNG logo for watermarking. Transparency recommended.")
    add_field(doc, "Primary/Secondary/Accent Color", "Color pickers", "Brand Kit tab", "Brand color palette in hex.")
    add_field(doc, "Font Family", "Dropdown", "Brand Kit tab", "Font for burned captions and text overlays.", "Inter, Montserrat, Poppins, Impact, etc.")
    add_field(doc, "Watermark Position", "Dropdown", "Brand Kit tab", "Where logo appears on videos.")
    add_field(doc, "Watermark Size", "Slider (5-40%)", "Brand Kit tab", "How large the watermark is. Default: 15%.")
    add_field(doc, "Save Brand Kit", "Button", "Brand Kit tab", "Saves settings. Applies to future uploads.")

    # Voice
    doc.add_page_break()
    add_heading(doc, "Voice Tab", level=2)
    add_field(doc, "Persona", "Text input", "Voice tab", "Describe yourself. AI uses this to match your voice.", "friendly dog mom who loves puns")
    add_field(doc, "Tone", "Text input", "Voice tab", "Emotional tone for AI content.", "warm, witty, enthusiastic")
    add_field(doc, "Signature Phrases", "Text input", "Voice tab", "Comma-separated phrases AI uses naturally.", "pawsome, fur-real, zoomies", "3-5 phrases is plenty.")
    add_field(doc, "Phrases to Avoid", "Text input", "Voice tab", "Words AI should never use.", "literally, basically")
    add_field(doc, "Reply Mode", "Dropdown", "Voice tab", "Suggest (manual approval) or Auto (auto-reply).", "Start with Suggest. Switch to Auto only when you trust the AI.")
    add_field(doc, "Reply Length", "Dropdown", "Voice tab", "Short (1 sentence), Medium (2-3), Long (3-5).", "Most platforms favor short replies.")

    # API Keys
    doc.add_page_break()
    add_heading(doc, "API Keys Tab", level=2)
    add_field(doc, "Get keys", "Link", "API Keys tab", "Opens platform's developer portal in new tab.")
    add_field(doc, "Client ID / Secret / API Token", "Password input", "API Keys tab", "Paste credentials. Encrypted before storage. Shows masked preview.", "Saved as 1234....com, never shown in full.", "Never share these values.")
    add_field(doc, "Save [Platform] keys", "Button", "API Keys tab", "Saves credentials. Button disabled until a field has a value.")
    add_field(doc, "Trash icon", "Button", "API Keys tab", "Deletes a saved credential.")

    # Assets
    doc.add_page_break()
    add_heading(doc, "Assets Tab", level=2)
    add_field(doc, "Watermark (PNG)", "File input", "Assets tab", "Upload PNG logo for video watermarking.")
    add_field(doc, "Intro Clip (MP4)", "File input", "Assets tab", "Upload video clip to prepend to videos.")
    add_field(doc, "Outro Clip (MP4)", "File input", "Assets tab", "Upload video clip to append to videos.")
    add_field(doc, "Background Music (MP3)", "File input", "Assets tab", "Upload music to mix under videos.")

    # Settings
    doc.add_page_break()
    add_heading(doc, "Settings Tab", level=2)
    add_field(doc, "Brand Handle", "Text input", "Settings tab", "Your social media handle. Used in AI captions.", "@yourdog", "Include the @ symbol.")
    add_field(doc, "Content Niche", "Text input", "Settings tab", "What kind of content you make. Affects ALL AI generation.", "dog / pet content or fitness and workouts", "Be specific. This single field affects everything.")
    add_field(doc, "Save Settings", "Button", "Settings tab", "Saves settings. Immediately affects all AI content.")

    # AI Agent
    doc.add_page_break()
    add_heading(doc, "AI Agent Chat", level=2)
    add_field(doc, "Chat Input", "Textarea", "Agent chat bottom", "Type any question or command. Press Enter to send.", "Type 'What should I post this week?' and press Enter.", "The agent can call up to 38 tools.")
    add_field(doc, "Send Button", "Button", "Agent chat bottom right", "Sends your message. Disabled while agent is thinking.")
    add_field(doc, "Suggested Questions", "Buttons", "Agent chat empty state", "4 pre-written questions. Click any to fill input.")
    add_field(doc, "Tool Call Indicators", "Display chips", "Agent chat response", "Shows which tools the agent called. Expandable.")
    add_field(doc, "New Conversation", "Button", "Agent chat header", "Starts fresh conversation. Previous one is saved.")

    add_para(doc, "")
    add_para(doc, "Ready-to-use Agent Commands:", bold=True, size=12)
    commands = [
        "What should I post this week?",
        "Give me 3 content ideas",
        "What's trending in my niche?",
        "Schedule my ready videos at optimal times",
        "Analyze this viral video: [paste YouTube URL]",
        "What analyzed scripts do I have?",
        "What frameworks have I saved?",
        "Show me my top performing videos",
        "Generate a thumbnail for my latest video",
        "Repurpose my latest video into 5 clips",
        "Translate my latest video to Spanish and French",
        "Generate 5 hooks for my latest video",
        "What are my competitors up to?",
        "Fetch new comments and suggest replies",
        "What is my brand name and primary color?",
        "Add @competitor_handle on YouTube as a competitor",
    ]
    for cmd in commands:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f'  "{cmd}"')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xE8, 0x6A, 0x00)

    # ============ PART 3: RECORDING CHECKLIST ============
    doc.add_page_break()
    add_heading(doc, "PART 3: RECORDING CHECKLIST", level=1, color=RGBColor(0xE8, 0x6A, 0x00))

    add_heading(doc, "Phase 1: Pre-Recording Setup (do once)", level=2)
    add_para(doc, "1.1 Deploy ContentForge to Vercel", bold=True, size=12)
    for item in [
        "Push code to GitHub repository",
        "Import repo into Vercel (vercel.com/new)",
        "Set ENCRYPTION_KEY env var (run: openssl rand -hex 32)",
        "Set CRON_SECRET env var (run: openssl rand -hex 16)",
        "Switch database to Postgres (Vercel Storage)",
        "Update prisma/schema.prisma: change sqlite to postgresql",
        "Redeploy and verify site loads at your Vercel URL",
    ]:
        p = doc.add_paragraph()
        run = p.add_run("  [  ]  ")
        run.font.size = Pt(12)
        run = p.add_run(item)
        run.font.size = Pt(11)

    add_para(doc, "1.2 Set Up Your Content Environment", bold=True, size=12)
    for item in [
        "Open Settings tab and set Brand Handle (e.g. @yourdog)",
        "Open Settings tab and set Content Niche (e.g. dog / pet content)",
        "Open Brand Kit tab and upload your logo (PNG with transparency)",
        "Open Brand Kit tab and set primary color, font, watermark position",
        "Open Voice tab and fill in Persona, Tone, Signature Phrases",
        "Open Assets tab and upload watermark, intro, outro, music",
        "Open API Keys tab and add at least YouTube credentials",
        "Open Social tab and connect at least one platform via OAuth",
        "Upload 2-3 real videos (so Library has content to show)",
        "Upload 1-2 photos as a slideshow (to demo photo-to-video)",
        "Wait for all videos to reach ready status",
        "Go to Ideas tab and click Generate 5 ideas",
        "Go to Trends tab and click Refresh trends",
        "Go to Analyzer tab and analyze 1-2 viral video URLs",
        "Go to Competitors tab and add 1-2 competitors and click Check now",
    ]:
        p = doc.add_paragraph()
        run = p.add_run("  [  ]  ")
        run.font.size = Pt(12)
        run = p.add_run(item)
        run.font.size = Pt(11)

    add_para(doc, "1.3 Prepare Recording Environment", bold=True, size=12)
    for item in [
        "Close all unnecessary browser tabs and applications",
        "Hide browser bookmarks bar (Ctrl+Shift+B in Chrome)",
        "Set browser to full screen / maximize window",
        "Set screen resolution to 1920x1080 or higher",
        "Enable screen magnifier tool",
        "Turn off notifications (Do Not Disturb mode)",
        "Test microphone: record 10 seconds and playback",
        "Position microphone 6-8 inches from mouth",
        "Have the video script document open on a second monitor",
        "Set browser zoom to 100% (Ctrl+0)",
    ]:
        p = doc.add_paragraph()
        run = p.add_run("  [  ]  ")
        run.font.size = Pt(12)
        run = p.add_run(item)
        run.font.size = Pt(11)

    add_heading(doc, "Phase 2: Recording Session (per video)", level=2)
    add_para(doc, "Before Hitting Record:", bold=True, size=12)
    for item in [
        "Read the video script segment from this document",
        "Navigate to the starting tab in ContentForge",
        "Prepare any data you will need (e.g. have a URL ready to paste)",
        "Clear any test data that should not be in the video",
        "Take a sip of water",
        "Deep breath, relax your shoulders",
        "Start screen recorder",
        "Wait 3 seconds before speaking (clean intro for editing)",
    ]:
        p = doc.add_paragraph()
        run = p.add_run("  [  ]  ")
        run.font.size = Pt(12)
        run = p.add_run(item)
        run.font.size = Pt(11)

    add_para(doc, "During Recording:", bold=True, size=12)
    for item in [
        "Speak clearly and at a moderate pace",
        "Move mouse slowly when pointing at UI elements",
        "Pause 2 seconds after clicking buttons (let UI respond)",
        "If you make a mistake, pause 3 seconds, then re-do the segment",
        "Zoom into important UI elements using screen magnifier",
        "Do not read the script word-for-word, paraphrase naturally",
        "End with the standard outro",
        "Wait 3 seconds after finishing before stopping recording",
    ]:
        p = doc.add_paragraph()
        run = p.add_run("  [  ]  ")
        run.font.size = Pt(12)
        run = p.add_run(item)
        run.font.size = Pt(11)

    add_heading(doc, "Phase 3: Post-Recording", level=2)
    add_para(doc, "Editing:", bold=True, size=12)
    for item in [
        "Import raw recordings into video editor (CapCut, DaVinci Resolve)",
        "Cut out mistakes and long pauses",
        "Add zoom effects on key UI elements",
        "Add intro/outro graphics (consistent across all videos)",
        "Add background music (low volume, -20dB below narration)",
        "Add captions/subtitles (use auto-caption, then fix errors)",
        "Export at 1080p MP4, H.264, 10-15 Mbps bitrate",
    ]:
        p = doc.add_paragraph()
        run = p.add_run("  [  ]  ")
        run.font.size = Pt(12)
        run = p.add_run(item)
        run.font.size = Pt(11)

    add_para(doc, "Publishing:", bold=True, size=12)
    for item in [
        "Create YouTube playlist: ContentForge Tutorial Series",
        "Upload video with title: ContentForge: [Feature Name] Tutorial",
        "Write description with timestamps + link to platform",
        "Add thumbnail (use the AI thumbnail generator)",
        "Add to playlist",
        "Add end screen pointing to next video in series",
    ]:
        p = doc.add_paragraph()
        run = p.add_run("  [  ]  ")
        run.font.size = Pt(12)
        run = p.add_run(item)
        run.font.size = Pt(11)

    # Equipment
    doc.add_page_break()
    add_heading(doc, "Equipment & Software", level=2)
    add_para(doc, "Screen Recording:", bold=True, size=12)
    add_para(doc, "- OBS Studio (FREE) - Best overall, Windows/Mac/Linux")
    add_para(doc, "- Loom (FREE tier) - Easiest, browser-based, 5 min limit")
    add_para(doc, "- QuickTime Player (FREE, Mac) - Simple, built into macOS")
    add_para(doc, "- Camtasia ($299) - Professional recording + editing")
    add_para(doc, "")
    add_para(doc, "Microphone:", bold=True, size=12)
    add_para(doc, "- Blue Yeti USB Mic (~$130) - Industry standard for YouTubers")
    add_para(doc, "- Rode NT-USB (~$170) - Higher quality, still USB-simple")
    add_para(doc, "- AirPods/headphones mic - Okay for starting out")
    add_para(doc, "")
    add_para(doc, "Video Editing:", bold=True, size=12)
    add_para(doc, "- CapCut (FREE) - Best free editor, desktop + mobile")
    add_para(doc, "- DaVinci Resolve (FREE) - Professional-grade, steep learning curve")
    add_para(doc, "- iMovie (FREE, Mac) - Simple, good for beginners")
    add_para(doc, "- Adobe Premiere Pro ($22/mo) - Industry standard")
    add_para(doc, "")
    add_para(doc, "For AI Video Generation (Higgsfield, etc.):", bold=True, size=12)
    add_para(doc, "- Use NARRATION text as the script input")
    add_para(doc, "- Describe VISUAL cues as scene directions")
    add_para(doc, "- Best for intro/marketing videos, not precise UI tutorials")
    add_para(doc, "- Screen recording is recommended for tutorial videos showing actual UI")

    # ============ OUTRO ============
    doc.add_page_break()
    add_heading(doc, "Outro Template (Use for All Videos)", level=1)
    add_narration(doc, "And that's it for this feature! If you found this helpful, hit like and subscribe for more ContentForge tutorials. Check the description for links to other videos in this series, and visit ContentForge dot com to start automating your content today. See you in the next video!")
    add_para(doc, "")
    add_para(doc, "Standard outro elements:", bold=True)
    add_para(doc, "- Subscribe button overlay")
    add_para(doc, "- Link to next video in playlist")
    add_para(doc, "- Link to ContentForge platform")
    add_para(doc, "- Social media handles")
    add_para(doc, "- Background music fade-out")

    # ============ SAVE ============
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")
    print(f"Total videos: 23")
    print(f"Total tabs covered: 20")
    print(f"Total runtime: ~85 minutes")

if __name__ == "__main__":
    build()
