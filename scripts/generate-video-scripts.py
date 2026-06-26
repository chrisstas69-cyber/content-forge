#!/usr/bin/env python3
"""
ContentForge — Complete Video Tutorial Script Generator
Creates a comprehensive .docx with video scripts for every feature.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

OUTPUT_PATH = "/home/z/my-project/download/ContentForge-Video-Scripts.docx"

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h

def add_para(doc, text, bold=False, italic=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_narration(doc, text):
    """Add narration script with distinct formatting"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"NARRATION: {text}")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p

def add_visual(doc, text):
    """Add visual cue with distinct formatting"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"VISUAL: {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run.bold = True
    return p

def add_tip(doc, text):
    """Add pro tip"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"PRO TIP: {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    run.bold = True
    return p

def add_video_segment(doc, number, title, duration, description):
    """Add a video segment header"""
    doc.add_page_break()
    add_heading(doc, f"Video {number}: {title}", level=1, color=RGBColor(0xE8, 0x6A, 0x00))
    add_para(doc, f"Duration: ~{duration} minutes", italic=True, size=10, color=RGBColor(0x66, 0x66, 0x66))
    add_para(doc, description, size=11, space_after=12)

def build_document():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ============ COVER PAGE ============
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ContentForge")
    run.bold = True
    run.font.size = Pt(42)
    run.font.color.rgb = RGBColor(0xE8, 0x6A, 0x00)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Complete Video Tutorial Script")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("A step-by-step guide for creating tutorial videos\ncovering every feature of the AI Content Automation Platform")
    run.font.size = Pt(13)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("21 Video Segments  |  Total Runtime: ~75 minutes\n\nFor use with screen recording, live narration, or AI video generation (Higgsfield, etc.)")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ============ TABLE OF CONTENTS ============
    doc.add_page_break()
    add_heading(doc, "Table of Contents", level=1)

    videos = [
        ("1", "Platform Overview & Welcome", "2 min"),
        ("2", "Dashboard Tour", "3 min"),
        ("3", "Uploading Videos", "5 min"),
        ("4", "Uploading Photos → Video", "4 min"),
        ("5", "Video Library Management", "3 min"),
        ("6", "AI Agent Chat", "5 min"),
        ("7", "Content Ideation Engine", "4 min"),
        ("8", "AI Generation (Thumbnails, Images, B-roll)", "5 min"),
        ("9", "Social Accounts & API Keys", "6 min"),
        ("10", "Scheduling & Calendar", "4 min"),
        ("11", "Trend Research", "3 min"),
        ("12", "Analytics Dashboard", "4 min"),
        ("13", "Insights Engine", "3 min"),
        ("14", "Auto-Comment Replies", "5 min"),
        ("15", "Voice Profile Setup", "3 min"),
        ("16", "Competitor Monitoring", "4 min"),
        ("17", "Brand Kit", "3 min"),
        ("18", "Content Repurposing", "4 min"),
        ("19", "Subtitle Translation", "3 min"),
        ("20", "Hook A/B Testing", "4 min"),
        ("21", "Multi-Format Publishing", "3 min"),
    ]

    for num, title, duration in videos:
        p = doc.add_paragraph()
        run = p.add_run(f"Video {num}: {title}")
        run.font.size = Pt(11)
        tab = p.add_run(f"\t{duration}")
        tab.font.size = Pt(10)
        tab.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ============ INTRODUCTION ============
    doc.add_page_break()
    add_heading(doc, "Introduction", level=1)

    add_para(doc, "This document contains complete video scripts for every feature of ContentForge, the AI-powered content automation platform. Each video segment is designed to be self-contained — you can record them in any order or combine multiple segments into longer videos.")

    add_para(doc, "How to use this script:", bold=True, size=12, space_after=4)
    add_para(doc, "1. Each video segment includes a NARRATION script (what to say), VISUAL cues (what to show on screen), and PRO TIPS.")
    add_para(doc, "2. For screen recordings: Follow the VISUAL cues while reading the NARRATION.")
    add_para(doc, "3. For AI video generation (Higgsfield, etc.): Use the NARRATION as the script input and describe the VISUAL cues as scene descriptions.")
    add_para(doc, "4. Total runtime is approximately 75 minutes if recorded as individual segments. You can combine related segments into longer compilation videos.")

    add_para(doc, "Recording tips:", bold=True, size=12, space_after=4)
    add_para(doc, "- Record in 1080p or 4K — viewers will pause to read UI text")
    add_para(doc, "- Use a clean desktop (hide bookmarks bar, close other tabs)")
    add_para(doc, "- Zoom in on important UI elements using screen magnifier")
    add_para(doc, "- Add captions/subtitles for accessibility")
    add_para(doc, "- Include a consistent intro/outro across all videos")

    # ============ VIDEO 1: PLATFORM OVERVIEW ============
    add_video_segment(doc, "1", "Platform Overview & Welcome", "2 min",
        "A high-level introduction to ContentForge — what it is, who it's for, and the complete content lifecycle it covers.")

    add_visual(doc, "Open browser to ContentForge dashboard. Slow zoom out to show the full interface. Hover over each nav tab briefly.")
    add_narration(doc, "Welcome to ContentForge — the all-in-one AI content automation platform that handles your entire content workflow from idea to published post. Whether you're a pet creator, fitness influencer, or business brand, ContentForge automates the repetitive work so you can focus on creating.")
    add_visual(doc, "Show the 18 nav tabs across the top. Highlight: Upload, Library, Ideas, Generate, Social, Analytics, Comments, Competitors.")
    add_narration(doc, "The platform covers the complete content lifecycle: ideation, creation, editing, publishing, analytics, and iteration. You can upload videos or photos, let AI edit them with captions and voiceover, schedule posts across five social platforms, monitor competitors, auto-reply to comments, and learn what works through built-in analytics.")
    add_visual(doc, "Click on the AI Agent floating button (bottom right). Show the chat opening.")
    add_narration(doc, "And the best part? There's a built-in AI agent that can do almost everything for you. Just ask it to generate ideas, schedule posts, check trends, or analyze your performance — and it executes. Let's dive into each feature.")
    add_tip(doc, "This video works great as a YouTube intro or a homepage hero video. Keep it under 2 minutes to maintain attention.")

    # ============ VIDEO 2: DASHBOARD TOUR ============
    add_video_segment(doc, "2", "Dashboard Tour", "3 min",
        "Walk through the main dashboard, explaining each stat card and quick action.")

    add_visual(doc, "Start on the Dashboard tab. Point to each stat card as you mention it.")
    add_narration(doc, "The Dashboard is your command center. At the top, you'll see six stat cards: Total Videos, Processing, Ready, Scheduled, Published, and Average Viral Score. These update in real-time as your content moves through the pipeline.")
    add_visual(doc, "Scroll down to Quick Actions section. Hover over each button.")
    add_narration(doc, "Below the stats, you'll find Quick Actions — shortcuts to the most common tasks: Upload new videos, Generate content ideas with AI, Generate AI thumbnails and B-roll, View library, and Content calendar. These are color-coded by priority — orange buttons are AI-powered features.")
    add_visual(doc, "Scroll to the 'How it works' section at the bottom.")
    add_narration(doc, "At the bottom, there's a quick reference showing the complete workflow: Upload, transcribe, edit, score, multi-format generation, and publish. This is great for new users to understand the full pipeline at a glance.")
    add_visual(doc, "Point to the floating orange sparkle button in the bottom right corner.")
    add_narration(doc, "And don't forget the AI Agent — this floating button is always available no matter which tab you're on. Click it anytime to ask questions or trigger actions without navigating through menus.")
    add_tip(doc, "Mention that the niche shown in the welcome message adapts to whatever the user set in Settings — it's not hardcoded to 'dog content'.")

    # ============ VIDEO 3: UPLOADING VIDEOS ============
    add_video_segment(doc, "3", "Uploading Videos", "5 min",
        "Complete walkthrough of the video upload process, including all editing options.")

    add_visual(doc, "Click on the Upload tab. Show the drag-and-drop zone.")
    add_narration(doc, "Let's upload our first video. Click the Upload tab. You can either drag and drop video files here, or click to browse. ContentForge supports MP4, MOV, and WebM — and you can upload multiple files at once.")
    add_visual(doc, "Select a video file. It appears in the file list below the dropzone.")
    add_narration(doc, "Once you've selected your files, they appear in this list. You can remove any file by clicking the X button. Now let's look at the editing options.")
    add_visual(doc, "Scroll down to 'Auto-Editing Options'. Point to each toggle and dropdown.")
    add_narration(doc, "The Auto-Editing Options let you control how AI processes your video. First, 'Burn AI captions' — this transcribes your audio and burns subtitles directly onto the video. Great for accessibility and engagement. Next, 'Auto-trim silence' — removes dead air from your video.")
    add_visual(doc, "Show the Watermark, Background Music, Intro Clip, and Outro Clip dropdowns.")
    add_narration(doc, "You can add a watermark — just upload a PNG with transparency in the Assets tab first. Same for background music, intro clips, and outro clips — upload them once in Assets, and they'll be available in this dropdown for every video.")
    add_visual(doc, "Scroll to the 'AI Voiceover' section. Toggle it on.")
    add_narration(doc, "Now here's a powerful feature — AI Voiceover. When enabled, the system generates a script based on your niche, then synthesizes it as audio using the voice you select. You can choose from seven voices, set the tone — like 'funny, energetic, engaging' — and decide whether to replace the original audio or mix the voiceover on top.")
    add_visual(doc, "Click the 'Upload' button. Show the video appearing in the Library with a 'pending' status, then transitioning to 'editing' with a progress bar.")
    add_narration(doc, "Once you click Upload, the video enters the processing pipeline. You'll see it appear in your Library with a progress bar showing each step: probing metadata, extracting audio, transcribing with AI, analyzing viral potential, editing with FFmpeg, and generating multi-format outputs. The whole process takes about 30 to 60 seconds for a typical short video.")
    add_tip(doc, "Mention that the viral score is generated by AI comparing the video to current trends — it's not random.")

    # ============ VIDEO 4: UPLOADING PHOTOS ============
    add_video_segment(doc, "4", "Uploading Photos → Video", "4 min",
        "How to upload still photos and have them converted into videos with voiceover, music, and captions.")

    add_visual(doc, "On the Upload tab, point to the purple 'Upload Photos Instead' button below the video dropzone.")
    add_narration(doc, "Did you know you can create videos from photos? ContentForge can turn your still images into engaging videos with Ken Burns zoom effects, AI voiceover, burned captions, and background music. Perfect for photo dumps, before-and-after content, or any time you don't have video footage.")
    add_visual(doc, "Click 'Upload Photos Instead'. Select 3-4 photos. They appear in a separate 'Photos → Slideshow' section.")
    add_narration(doc, "Click 'Upload Photos Instead' and select one or more images. They appear in this purple section, separate from your video uploads. If you upload multiple photos, they'll be combined into a slideshow with crossfade transitions.")
    add_visual(doc, "Show the photo slideshow settings: seconds per photo, transition speed, and custom voiceover script field.")
    add_narration(doc, "Now you have photo-specific settings. 'Seconds per photo' controls how long each image displays — default is 5 seconds. 'Transition' controls the crossfade duration. And here's the best part — you can write a custom voiceover script, or leave it blank and let AI generate one based on your niche and current trends.")
    add_visual(doc, "Type a custom script in the voiceover script field.")
    add_narration(doc, "If you provide a script, that exact text will be used for the voiceover AND the burned captions. This gives you full creative control. If you leave it blank, AI writes a script tailored to your content niche and brand voice.")
    add_visual(doc, "Click Upload. Show the video appearing in Library with status 'editing'.")
    add_narration(doc, "Click Upload, and the system converts your photos into a video, generates the voiceover, burns captions, adds music, and produces all three formats. The result appears in your Library looking just like a regular video — ready to publish.")
    add_tip(doc, "Show an example of a finished photo-to-video result. The AI-generated title and viral score make it clear this is a complete, publish-ready piece of content.")

    # ============ VIDEO 5: VIDEO LIBRARY ============
    add_video_segment(doc, "5", "Video Library Management", "3 min",
        "How to manage processed videos — publish, repurpose, translate, delete, and understand status badges.")

    add_visual(doc, "Navigate to the Library tab. Show the grid of video cards.")
    add_narration(doc, "The Library is where all your processed videos live. Each card shows the thumbnail, viral score, status, AI-generated title, duration, file size, and hashtags. Let's break down what each element means.")
    add_visual(doc, "Point to the viral score circle (top right of a card).")
    add_narration(doc, "This circle shows your AI-generated viral score from 0 to 100. Green means 80 or above — high viral potential. Amber means 50 to 79 — decent. Red means below 50 — might need improvement. The score is based on hook strength, emotional pull, trend alignment, and rewatchability.")
    add_visual(doc, "Point to the status badge (top left). Show different statuses: pending, editing, ready, published, failed.")
    add_narration(doc, "The status badge shows where the video is in the pipeline. 'Pending' means it's queued. 'Editing' means FFmpeg is processing. 'Ready' means it's done and ready to publish. 'Published' means it's been pushed to social platforms. 'Failed' means something went wrong — click Retry to reprocess.")
    add_visual(doc, "Point to the format badges below the video preview.")
    add_narration(doc, "These green badges show which formats have been generated: 16:9 for YouTube and Facebook, 9:16 for TikTok and Reels, and 1:1 for Instagram feed. Every video automatically gets all three formats.")
    add_visual(doc, "Click the Publish button on a ready video. Show the publish dialog briefly.")
    add_narration(doc, "When a video is ready, you'll see three action buttons. 'Publish' opens the publishing dialog where you select platforms and schedule. 'Translate' lets you translate the caption to 12 languages. And 'Repurpose' appears on videos longer than 90 seconds — it uses AI to cut the video into short clips.")
    add_visual(doc, "Click the Repurpose button. Show the confirmation message.")
    add_narration(doc, "Clicking Repurpose triggers AI to analyze the transcript and identify the most engaging moments. It then extracts those clips as separate videos, each with their own titles, captions, and viral scores. One long video becomes five short clips — a massive time saver.")
    add_tip(doc, "Mention that clips show a purple 'Clip' badge and display their source timestamp — 'from 45s' means it was cut from the 45-second mark of the parent video.")

    # ============ VIDEO 6: AI AGENT CHAT ============
    add_video_segment(doc, "6", "AI Agent Chat", "5 min",
        "Deep dive into the AI Agent — the floating chat that can see your data and take actions.")

    add_visual(doc, "Click the floating orange sparkle button in the bottom right. The chat sidebar opens.")
    add_narration(doc, "This is the ContentForge AI Agent — your personal content strategist, social media manager, and analyst, all in one. It's not just a chatbot. It has access to 33 tools that let it see your data and take real actions in your account.")
    add_visual(doc, "Show the welcome screen with suggested questions.")
    add_narration(doc, "When you first open it, you'll see suggested questions to get started. Let's try one.")
    add_visual(doc, "Click 'What should I post this week?' Show the agent calling tools (list_videos, get_settings, search_trends) and then generating a response.")
    add_narration(doc, "Watch what happens when I ask 'What should I post this week?' The agent calls three tools in sequence: it lists your videos to see what you have, reads your settings to know your niche, and searches current trends. Then it synthesizes all that data into a personalized weekly strategy.")
    add_visual(doc, "Show the response with tool call indicators ('used tool: list_videos' etc.).")
    add_narration(doc, "These little chips show which tools the agent used. You can expand them to see the actual data returned. This transparency builds trust — you know exactly what the agent based its advice on.")
    add_visual(doc, "Type a new question: 'Schedule my ready videos at optimal times'")
    add_narration(doc, "The agent can also take actions. If I ask it to schedule my ready videos, it will actually create scheduled posts at optimal times for each platform. It's not just giving advice — it's executing your workflow.")
    add_visual(doc, "Show the agent confirming the action was taken.")
    add_narration(doc, "After taking an action, the agent confirms what it did. You can always verify by checking the Scheduled tab.")
    add_visual(doc, "Click 'New conversation' to start fresh.")
    add_narration(doc, "Conversations are saved, so you can reference them later. Click the new conversation button to start fresh, or continue an existing conversation to maintain context.")
    add_tip(doc, "Emphasize that the agent remembers context within a conversation. If you ask 'what's my brand handle?' it already knows from previous messages.")

    # ============ VIDEO 7: CONTENT IDEATION ============
    add_video_segment(doc, "7", "Content Ideation Engine", "4 min",
        "How to use the Ideas tab to generate AI content ideas based on trends, analytics, and performance insights.")

    add_visual(doc, "Navigate to the Ideas tab. Show the empty state with the 'Generate 5 ideas' button.")
    add_narration(doc, "Never know what to post? The Ideas tab uses AI to generate content ideas tailored to your niche, grounded in current trends and your actual performance data. Let's generate some ideas.")
    add_visual(doc, "Click 'Generate 5 ideas'. Show the loading spinner, then the generated ideas appearing.")
    add_narration(doc, "Click 'Generate 5 ideas' and the system pulls together your current trends, analytics insights, recent videos, and niche settings. It sends all that context to the AI, which generates five original content ideas. This takes about 15 to 20 seconds.")
    add_visual(doc, "Scroll through the generated ideas. Point to each element: title, viral score, concept, script outline, format, source.")
    add_narration(doc, "Each idea card includes: a scroll-stopping title, a predicted viral score from 0 to 100, a concept description, a full script outline with Hook, Body, and CTA sections, the recommended format — like 9:16 for TikTok — and the source, which shows what data informed this idea.")
    add_visual(doc, "Expand the script outline section of one idea.")
    add_narration(doc, "The script outline is production-ready. It gives you a hook for the first three seconds, the body content, and a call to action. You can take this outline and shoot your video immediately.")
    add_visual(doc, "Click 'Dismiss' on one idea to remove it.")
    add_narration(doc, "If you don't like an idea, click Dismiss to remove it. Used ideas move to a separate section so you can track what you've already created. The agent can also generate ideas — just ask it 'give me 3 content ideas' in the chat.")
    add_tip(doc, "Mention that ideas get smarter over time. As you publish more and the Insights engine learns what works for you, ideation incorporates those patterns.")

    # ============ VIDEO 8: AI GENERATION ============
    add_video_segment(doc, "8", "AI Generation (Thumbnails, Images, B-roll)", "5 min",
        "Complete guide to the Generate tab — AI thumbnails with image upload, AI images, and AI B-roll video generation.")

    add_visual(doc, "Navigate to the Generate tab. Show the three sub-tabs: AI Thumbnail, AI Image, AI B-roll Video.")
    add_narration(doc, "The Generate tab is where AI creates visual assets for you. There are three modes: AI Thumbnails for eye-catching video thumbnails, AI Images for any image from text, and AI B-roll Video for generating video clips from text prompts. Let's start with thumbnails.")
    add_visual(doc, "With 'AI Thumbnail' selected, enter a video title. Show the 'Upload your own image' section.")
    add_narration(doc, "For thumbnails, enter your video title. Now here's the exciting part — you can upload your own photo and AI will transform it into a stylized thumbnail. Click 'Click to upload your photo' and select an image — like a photo of your dog.")
    add_visual(doc, "Upload a photo. Show the preview appearing, plus the style strength slider.")
    add_narration(doc, "Once uploaded, you'll see a preview and a style strength slider. At 10 to 30 percent, your photo stays very recognizable with minimal AI styling. At 30 to 50 percent, you get a balanced mix of your photo and thumbnail styling. Above 50 percent, the AI transforms the image more dramatically. The default 35 percent is a great starting point.")
    add_visual(doc, "Click 'Generate from your photo'. Show the loading state, then the result appearing in the grid below.")
    add_narration(doc, "Click 'Generate from your photo' and the system uses Replicate's SDXL image-to-image model to transform your upload. This takes about 30 to 60 seconds. The result appears in the Recent Generations grid below. Without an upload, AI generates from scratch using the built-in image generator — no API key needed.")
    add_visual(doc, "Switch to the 'AI Image' sub-tab. Enter a prompt and generate.")
    add_narration(doc, "The AI Image mode generates any image from a text prompt. Just describe what you want — like 'a cozy living room with a sleeping cat, warm lighting' — and click Generate. This uses the built-in AI, so no API key required.")
    add_visual(doc, "Switch to the 'AI B-roll Video' sub-tab. Show the warning about Replicate API token.")
    add_narration(doc, "AI B-roll Video generates actual video clips from text prompts — like 'cinematic shot of a golden retriever running on a beach at sunset.' This requires a Replicate API token, which you can add in the API Keys tab. Generation takes 2 to 5 minutes since it's creating actual video frames.")
    add_tip(doc, "Show the model name and status badge on each generated asset. Explain that B-roll videos show a loading spinner until ready.")

    # ============ VIDEO 9: SOCIAL ACCOUNTS & API KEYS ============
    add_video_segment(doc, "9", "Social Accounts & API Keys", "6 min",
        "How to add API credentials and connect social media accounts via OAuth.")

    add_visual(doc, "Navigate to the API Keys tab. Show the list of platforms: YouTube, TikTok, Instagram+Facebook, X, Replicate.")
    add_narration(doc, "Before you can publish to social platforms, you need to add API credentials. Go to the API Keys tab. You'll see five platform groups: YouTube, TikTok, Instagram plus Facebook, X, and Replicate for AI video generation. Each shows whether it's configured or not.")
    add_visual(doc, "Click 'Get keys' next to YouTube. Explain the Google Cloud Console process.")
    add_narration(doc, "Click 'Get keys' to open the platform's developer portal. For YouTube, you'll create a Google Cloud Project, enable the YouTube Data API, and create OAuth credentials. The system shows you the exact redirect URL to use — copy it and paste it into your Google Cloud Console.")
    add_visual(doc, "Show pasting a Client ID and Client Secret into the fields. Click 'Save YouTube keys'.")
    add_narration(doc, "Once you have your Client ID and Secret from Google, paste them here and click 'Save YouTube keys'. The values are encrypted before being stored in the database. You'll see a masked preview like '1234 dots dots dots dot com' — your secret is never displayed in full.")
    add_visual(doc, "Show the 'Configured' badge appearing on the YouTube card.")
    add_narration(doc, "Once saved, the YouTube card shows a green 'Configured' badge and the counter at the top updates — '1 out of 5 platforms configured.'")
    add_visual(doc, "Navigate to the Social tab. Show YouTube now showing 'API Ready' with a 'Connect' button.")
    add_narration(doc, "Now go to the Social tab. YouTube shows 'API Ready' with a Connect button. Click it to start the OAuth flow.")
    add_visual(doc, "Click 'Connect'. Show the redirect to Google's login page (mention this, don't show actual credentials).")
    add_narration(doc, "You'll be redirected to Google's login page. Authorize the app, and you'll be sent back to ContentForge. Your YouTube account now shows as connected with your handle and display name.")
    add_visual(doc, "Show the connected account with handle and display name. Point to the 'Needs Keys' badges on other platforms.")
    add_narration(doc, "Repeat this process for each platform you want to publish to. TikTok, Instagram, Facebook, and X each have their own developer portals — links are provided in the API Keys tab. The setup instructions in DEPLOYMENT.md walk you through each one step by step.")
    add_tip(doc, "Mention that Instagram and Facebook require videos hosted on a public URL for publishing — the user will need S3 or Cloudinary storage for those two platforms.")

    # ============ VIDEO 10: SCHEDULING & CALENDAR ============
    add_video_segment(doc, "10", "Scheduling & Calendar", "4 min",
        "How to schedule posts at optimal times and view them in the content calendar.")

    add_visual(doc, "From the Library, click 'Publish' on a ready video. Show the publish dialog with the 'When to publish' section.")
    add_narration(doc, "When you publish a video, you have three scheduling options. Let me show you. Click Publish on any ready video, and scroll down to the 'When to publish' section.")
    add_visual(doc, "Point to the three buttons: Publish now, Optimal times, Pick time.")
    add_narration(doc, "You can publish immediately, schedule at AI-determined optimal times, or pick a specific time. Let me explain each.")
    add_visual(doc, "Click 'Optimal times'. Show the explanation text that appears.")
    add_narration(doc, "Optimal times uses industry-standard peak engagement hours for each platform. YouTube posts at 3 PM Eastern, TikTok at 9 AM, noon, or 7 PM, Instagram at 11 AM, 2 PM, or 7 PM. Each platform gets its own optimal slot — so if you publish to all five, they'll spread across the day.")
    add_visual(doc, "Click 'Pick time'. Show the datetime picker.")
    add_narration(doc, "Pick time lets you choose a specific date and time. Useful for coordinated launches or holiday content.")
    add_visual(doc, "Select 'Optimal times' and click 'Schedule optimal'. Show the success message.")
    add_narration(doc, "Once scheduled, the post appears in the Scheduled tab with a blue 'Scheduled' badge showing the date and time. Vercel Cron checks every minute, and when the time comes, it automatically publishes to the selected platform.")
    add_visual(doc, "Navigate to the Calendar tab. Show the month grid with scheduled posts.")
    add_narration(doc, "The Calendar tab gives you a visual month view of all scheduled and published posts. Each day shows colored dots for each post — blue for scheduled, green for published, red for failed. Click the month navigation to look ahead or review past performance.")
    add_tip(doc, "Mention that the agent can also schedule posts — just ask it to 'schedule my ready videos at optimal times' and it handles everything.")

    # ============ VIDEO 11: TREND RESEARCH ============
    add_video_segment(doc, "11", "Trend Research", "3 min",
        "How to use the Trends tab to discover what's currently trending in your niche.")

    add_visual(doc, "Navigate to the Trends tab. Show the empty state with the 'Refresh trends' button.")
    add_narration(doc, "The Trends tab shows you what's currently trending in your niche across TikTok, Instagram, YouTube, and X. Trends are refreshed daily by an automated cron job, but you can also trigger a refresh manually.")
    add_visual(doc, "Click 'Refresh trends'. Show the loading state, then results appearing.")
    add_narration(doc, "Click 'Refresh trends' and the system performs web searches for trending content in your niche, then uses AI to extract and score the trends. This takes about 30 to 60 seconds since it's doing real web searches.")
    add_visual(doc, "Show the trends list. Point to each element: platform, type icon, content, summary, score, date.")
    add_narration(doc, "Each trend shows the platform, a type icon — hashtag, sound, format, or topic — the trend content, a summary explaining why it's trending, a relevance score from 0 to 100, and when it was discovered. Higher scores mean more relevant to your niche.")
    add_visual(doc, "Click a platform filter button (e.g., 'TikTok (10)').")
    add_narration(doc, "You can filter by platform using these buttons. This helps when you're planning platform-specific content.")
    add_visual(doc, "Point to the niche label at the top.")
    add_narration(doc, "The trends are specific to your niche — which you set in Settings. If you change your niche from 'dog content' to 'fitness content,' the next trend refresh will pull fitness trends instead. Trend data also feeds into the ideation engine, so generated ideas are grounded in what's actually trending right now.")
    add_tip(doc, "Explain that trends feed automatically into video analysis — when a video is processed, the AI includes current trending hashtags in its suggestions.")

    # ============ VIDEO 12: ANALYTICS DASHBOARD ============
    add_video_segment(doc, "12", "Analytics Dashboard", "4 min",
        "How to view cross-platform analytics and understand your content performance.")

    add_visual(doc, "Navigate to the Analytics tab. Show the four stat cards: Views, Likes, Comments, Shares.")
    add_narration(doc, "The Analytics tab pulls performance metrics from all your connected social platforms into one unified view. At the top, you'll see four stat cards: Total Views, Total Likes, Comments, and Shares across all platforms.")
    add_visual(doc, "Scroll down to the Per-Platform Breakdown.")
    add_narration(doc, "Below the totals, you'll see a per-platform breakdown showing how many posts and views each platform has generated. This helps you identify which platforms are performing best for your content.")
    add_visual(doc, "Scroll to the Top Performing Videos list.")
    add_narration(doc, "The Top Performing Videos section ranks your content by views. Each entry shows the rank, thumbnail, title, and key metrics. This is gold for understanding what resonates with your audience — double down on what works.")
    add_visual(doc, "Click 'Refresh metrics'.")
    add_narration(doc, "Metrics are refreshed automatically every hour by a cron job, but you can click 'Refresh metrics' to pull fresh data immediately. The system fetches view counts, likes, comments, and shares from each platform's API.")
    add_visual(doc, "Show the 'No analytics yet' empty state (mention what it looks like).")
    add_narration(doc, "If you haven't published any videos yet, you'll see an empty state prompting you to publish first. Once you have published posts, the analytics will populate within an hour.")
    add_tip(doc, "Mention that analytics data feeds into the Insights engine, which learns patterns from your performance and feeds them back into ideation.")

    # ============ VIDEO 13: INSIGHTS ENGINE ============
    add_video_segment(doc, "13", "Insights Engine", "3 min",
        "How the learning loop works — AI analyzes your analytics to generate actionable insights.")

    add_visual(doc, "Navigate to the Insights tab. Show the 'Refresh insights' button.")
    add_narration(doc, "The Insights tab is where the system learns from your performance. AI analyzes your last 30 videos and analytics data to identify patterns, opportunities, underperformers, and strategic recommendations.")
    add_visual(doc, "Click 'Refresh insights'. Show the loading state, then insights appearing.")
    add_narration(doc, "Click 'Refresh insights' and the system aggregates your video data, analytics metrics, hashtag performance, and current trends. It sends all this to the AI, which generates 5 to 8 actionable insights. This takes about 20 seconds.")
    add_visual(doc, "Show the insights list. Point to each type: Pattern, Opportunity, Underperformer, Recommendation.")
    add_narration(doc, "Each insight is color-coded by type. Green Patterns show what's working — like 'your puppy videos get 3x more views than training videos.' Blue Opportunities highlight untapped potential. Red Underperformers flag what's failing. Amber Recommendations give strategic advice.")
    add_visual(doc, "Expand the data section on one insight.")
    add_narration(doc, "Each insight includes supporting data — actual numbers from your analytics. Click 'view data' to see the metrics behind the insight. The confidence score shows how certain the AI is about the pattern.")
    add_visual(doc, "Point to the 'Feeds into ideation' text (or mention it).")
    add_narration(doc, "Insights feed directly into the ideation engine. When you generate content ideas, the AI references these insights to suggest content that aligns with what's actually working for you. It's a true learning loop — publish, analyze, learn, improve.")
    add_tip(doc, "Mention that insights are refreshed daily by a cron job, so they stay current as you publish more content.")

    # ============ VIDEO 14: AUTO-COMMENT REPLIES ============
    add_video_segment(doc, "14", "Auto-Comment Replies", "5 min",
        "How to fetch comments, review AI-suggested replies, and approve or edit them before posting.")

    add_visual(doc, "Navigate to the Comments tab. Show the empty state with 'Fetch new comments' button.")
    add_narration(doc, "The Comments tab is your comment management center. AI fetches comments from your published posts and generates replies in your voice. Let me show you how it works.")
    add_visual(doc, "Click 'Fetch new comments'. Show the loading state, then comments appearing.")
    add_narration(doc, "Click 'Fetch new comments' and the system pulls recent comments from YouTube, Instagram, Facebook, and X. For each new comment, AI generates a suggested reply based on your Voice Profile. This takes about 30 seconds depending on how many comments you have.")
    add_visual(doc, "Show a comment card with the AI suggested reply in the blue box.")
    add_narration(doc, "Each comment card shows the author name, platform, comment text, and the video it was posted on. Below that, in the blue box, is the AI-suggested reply — written in your voice, matching your persona and tone.")
    add_visual(doc, "Point to the three action buttons: Approve & Post, Edit, Ignore.")
    add_narration(doc, "You have three options. 'Approve and Post' sends the AI reply immediately. 'Edit' lets you modify the reply before posting — useful if you want to add a personal touch. 'Ignore' marks the comment as skipped.")
    add_visual(doc, "Click 'Edit' on a comment. Show the edit textarea.")
    add_narration(doc, "When you click Edit, the suggested reply becomes editable. Make your changes and click Post. The reply is published to the platform directly from ContentForge — no need to open each app separately.")
    add_visual(doc, "Show the filter buttons: Pending, Replied, Ignored, All.")
    add_narration(doc, "Use the filter tabs to view comments by status. 'Pending' shows comments awaiting your review. 'Replied' shows completed replies with timestamps. 'Ignored' shows comments you skipped.")
    add_visual(doc, "Mention the auto-reply mode (covered in Voice Profile video).")
    add_narration(doc, "If you enable auto-reply mode in the Voice tab, AI will post replies automatically without your approval — great for high-volume accounts. But most creators prefer the suggest mode for more control.")
    add_tip(doc, "Mention that comments are checked 4 times daily by a cron job, so you never miss new engagement.")

    # ============ VIDEO 15: VOICE PROFILE ============
    add_video_segment(doc, "15", "Voice Profile Setup", "3 min",
        "How to configure your AI voice profile for consistent replies and scripts.")

    add_visual(doc, "Navigate to the Voice tab. Show the voice profile form.")
    add_narration(doc, "The Voice tab is where you define your AI persona — the voice that generates comment replies, voiceover scripts, and other AI-written content. Getting this right makes everything sound like YOU, not a generic AI.")
    add_visual(doc, "Point to the Persona field.")
    add_narration(doc, "Start with the Persona field. Describe yourself in a sentence — like 'friendly dog mom who loves puns' or 'energetic fitness coach who motivates beginners.' This gives the AI context for who you are.")
    add_visual(doc, "Point to the Tone field.")
    add_narration(doc, "Next, set your Tone — like 'warm, witty, enthusiastic' or 'professional, authoritative, encouraging.' This controls the emotional register of AI-generated content.")
    add_visual(doc, "Point to the Signature Phrases field.")
    add_narration(doc, "Signature Phrases are words or expressions you use frequently — like 'pawsome, fur-real, zoomies' for a dog account. AI will weave these into replies and scripts naturally. Don't overdo it — 3 to 5 phrases is plenty.")
    add_visual(doc, "Point to the Phrases to Avoid field.")
    add_narration(doc, "Phrases to Avoid lets you blacklist words you never want AI to use. Maybe you hate 'literally' or 'basically' — add them here and AI will never use them.")
    add_visual(doc, "Point to the Reply Mode dropdown.")
    add_narration(doc, "Reply Mode has two options. 'Suggest' generates replies for your manual approval — recommended for most creators. 'Auto' posts replies immediately without approval — use this only if you trust the AI and have high comment volume.")
    add_visual(doc, "Point to the Reply Length dropdown.")
    add_narration(doc, "Reply Length controls how long AI replies are. Short is one sentence, Medium is two to three sentences, Long is three to five. Most social platforms favor short, punchy replies.")
    add_visual(doc, "Click 'Save Voice Profile'.")
    add_narration(doc, "Click Save, and your voice profile is stored. It immediately affects all AI-generated content — comment replies, voiceover scripts, and more. Update it anytime as your brand evolves.")
    add_tip(doc, "Encourage users to experiment with their voice profile. Try different personas and see which generates the best replies.")

    # ============ VIDEO 16: COMPETITOR MONITORING ============
    add_video_segment(doc, "16", "Competitor Monitoring", "4 min",
        "How to add competitors, get viral alerts, and use AI suggestions to create your own version.")

    add_visual(doc, "Navigate to the Competitors tab. Show the add competitor form.")
    add_narration(doc, "The Competitors tab lets you track other creators in your niche. When their posts go viral, you get an alert with AI suggestions on how to create your own version. Let's set one up.")
    add_visual(doc, "Select 'youtube' from the platform dropdown. Type a handle like 'mrbeast'. Click 'Add'.")
    add_narration(doc, "Select the platform — YouTube, Instagram, TikTok, or X — and type the competitor's handle without the @ symbol. Click Add and they're added to your tracked list.")
    add_visual(doc, "Show the competitor appearing in the 'Tracked competitors' list.")
    add_narration(doc, "The competitor now appears in your tracked list. The system will check their recent posts daily and alert you when something goes viral.")
    add_visual(doc, "Click 'Check now'. Show the loading state, then results.")
    add_narration(doc, "Click 'Check now' to fetch their latest posts immediately. AI analyzes each post, estimates a viral score from 0 to 100, and flags posts scoring 70 or above as viral alerts.")
    add_visual(doc, "Show the 'Viral alerts' section with alert cards.")
    add_narration(doc, "Viral alerts appear in this orange section. Each alert shows the post title, the competitor, the viral score, and an AI suggestion — like 'Create a similar video about X but add Y twist. Use hashtag Z to ride the trend.' This turns competitor research into actionable content ideas.")
    add_visual(doc, "Click 'View →' on an alert to open the original post.")
    add_narration(doc, "Click 'View' to open the competitor's original post in a new tab. Study what made it work, then use the AI suggestion to create your own version. You can also ask the agent 'what are my competitors up to?' for a quick summary.")
    add_tip(doc, "Suggest adding 3-5 competitors for good coverage. Too many will create alert fatigue.")

    # ============ VIDEO 17: BRAND KIT ============
    add_video_segment(doc, "17", "Brand Kit", "3 min",
        "How to set up your brand logo, colors, and fonts for consistent video branding.")

    add_visual(doc, "Navigate to the Brand Kit tab. Show the logo upload section.")
    add_narration(doc, "The Brand Kit tab stores your visual identity — logo, colors, and fonts — so every video gets consistent branding automatically. Let's set it up.")
    add_visual(doc, "Click 'Upload logo'. Select a PNG file. Show the logo preview appearing.")
    add_narration(doc, "Start by uploading your logo. PNG with transparency works best — it'll be applied as a watermark to every video. Click 'Upload logo' and select your file. The preview appears immediately.")
    add_visual(doc, "Scroll to the Brand Identity section. Show the color pickers.")
    add_narration(doc, "Next, set your brand colors. The Primary Color is your main brand color — used for accents and highlights. Secondary and Accent colors round out your palette. These aren't just for show — they're available to the AI when generating thumbnails and other assets.")
    add_visual(doc, "Show the Font Family dropdown.")
    add_narration(doc, "Choose your Font Family from the list. This affects burned-in captions and any text overlays. Inter, Montserrat, and Poppins are popular modern choices. Impact is great for thumbnail-style text.")
    add_visual(doc, "Show the Watermark Position and Size controls.")
    add_narration(doc, "Control where your logo appears — top left, top right, bottom left, bottom right, or center. The size slider controls how large the watermark is relative to the video — 15 percent is a good default, but adjust based on your logo's design.")
    add_visual(doc, "Click 'Save Brand Kit'.")
    add_narration(doc, "Click 'Save Brand Kit' and your settings are stored. Every video you upload from now on will automatically use your logo as a watermark in the position and size you specified. Update your brand kit anytime — changes apply to future uploads only.")
    add_tip(doc, "Mention that the agent can also update brand kit settings — just ask it to 'set my primary color to blue' or 'change my watermark to top right'.")

    # ============ VIDEO 18: CONTENT REPURPOSING ============
    add_video_segment(doc, "18", "Content Repurposing", "4 min",
        "How to turn one long video into multiple short clips automatically.")

    add_visual(doc, "Navigate to the Library. Find a video longer than 90 seconds. Point to the purple 'Repurpose' button.")
    add_narration(doc, "Content repurposing is one of the most powerful features in ContentForge. If you have a long video — say a 10-minute YouTube video — you can turn it into 5 short clips for TikTok, Reels, and Shorts with one click. Let me show you.")
    add_visual(doc, "Click 'Repurpose' on a long video. Show the confirmation message.")
    add_narration(doc, "Find any video longer than 90 seconds in your Library and click the purple 'Repurpose' button. The system confirms that repurposing has started.")
    add_visual(doc, "Explain what's happening behind the scenes (use text overlay or narration).")
    add_narration(doc, "Behind the scenes, AI analyzes the full transcript and identifies the 5 most engaging moments — based on hook strength, emotional peaks, and standalone value. Then FFmpeg extracts each clip as a separate video, and each clip runs through the full processing pipeline: captions, voiceover, multi-format generation, and viral scoring.")
    add_visual(doc, "Wait a moment, then refresh the Library. Show the new clip cards appearing with purple 'Clip' badges.")
    add_narration(doc, "After a few minutes, refresh your Library. You'll see new video cards with a purple 'Clip' badge. Each clip shows its own title, viral score, and the timestamp it was cut from — like 'from 45 seconds' means it starts at the 45-second mark of the original video.")
    add_visual(doc, "Point to the clip's title and score. Show that it's a fully processed video.")
    add_narration(doc, "Each clip is a fully processed video — it has its own AI-generated title, captions, hashtags, viral score, and all three formats. You can publish it just like any other video. One long video becomes five ready-to-publish short clips — a massive content multiplier.")
    add_visual(doc, "Show the agent chat. Ask it to repurpose a video.")
    add_narration(doc, "You can also trigger repurposing through the AI Agent. Just ask it to 'repurpose my latest video into 5 clips' and it handles everything. The agent will confirm when it's started and you can check back in a few minutes.")
    add_tip(doc, "Mention that repurposing works best on videos with clear speech — the AI needs the transcript to identify engaging moments.")

    # ============ VIDEO 19: SUBTITLE TRANSLATION ============
    add_video_segment(doc, "19", "Subtitle Translation", "3 min",
        "How to translate your video captions into 12 languages for international audiences.")

    add_visual(doc, "In the Library, click 'Translate' on a ready video. Show the Translate dialog.")
    add_narration(doc, "If you want to reach international audiences, ContentForge can translate your captions into 12 languages with one click. Let me show you.")
    add_visual(doc, "Show the language grid: Spanish, French, German, Italian, Portuguese, Dutch, Russian, Japanese, Korean, Chinese, Arabic, Hindi.")
    add_narration(doc, "The Translate dialog shows 12 supported languages. Each language shows a green checkmark if a translation already exists — so you can see at a glance what's been translated.")
    add_visual(doc, "Select 3 languages (e.g., Spanish, French, Japanese). Click 'Translate to 3'.")
    add_narration(doc, "Select the languages you want — you can pick multiple at once. Then click 'Translate to N' and AI translates your caption, description, and title into each selected language. This takes about 10 seconds per language.")
    add_visual(doc, "Show the translations appearing in the 'Existing translations' section.")
    add_narration(doc, "Translated text appears in the 'Existing translations' section. Each translation preserves your tone, emojis, and brand voice — it's not a literal translation, it's culturally adapted.")
    add_visual(doc, "Explain how translations are used (mention publishing to international audiences).")
    add_narration(doc, "Translations are stored with the video record. When you publish to platforms that support multiple languages — like YouTube — you can copy the translated text into the platform's language fields. This helps your content appear in international search results and reach viewers who don't speak English.")
    add_tip(doc, "Mention that the agent can also translate — just ask it to 'translate my latest video to Spanish and French'.")

    # ============ VIDEO 20: HOOK A/B TESTING ============
    add_video_segment(doc, "20", "Hook A/B Testing", "4 min",
        "How to generate hook variants and learn which opening styles work best for your audience.")

    add_visual(doc, "Explain the concept of hooks (use text overlay or narration).")
    add_narration(doc, "The first 3 seconds of a video decide whether someone scrolls past or watches to the end. That opening is called a 'hook.' ContentForge can generate 5 different hook variants for any video — each using a different psychological technique — so you can test which style works best for your audience.")
    add_visual(doc, "Show the API call or agent command to generate hooks. (Mention this is accessible via the agent.)")
    add_narration(doc, "To generate hooks, ask the AI Agent: 'generate 5 hooks for my latest video.' The agent calls the hook generation tool, which analyzes your video's transcript and creates 5 variants.")
    add_visual(doc, "Show the 5 hook styles: Question, Bold Claim, Curiosity Gap, Stat, Story.")
    add_narration(doc, "Each hook uses a different style. Question hooks ask something that makes viewers need the answer. Bold Claims challenge beliefs. Curiosity Gaps tease without revealing. Stats shock with numbers. Story hooks open with 'here's what happened when…'")
    add_visual(doc, "Show the hooks with their predicted scores and reasoning.")
    add_narration(doc, "Each hook includes a predicted viral score and reasoning — explaining why it works. Use the highest-scoring hook as the text overlay in your video's opening 3 seconds.")
    add_visual(doc, "Explain the learning loop (mention this happens automatically over time).")
    add_narration(doc, "Here's where it gets powerful. After your posts have been live for 24 hours, the system checks which hook style performed best based on actual engagement metrics. Over time, it learns which styles work for YOUR audience — and feeds that insight back into future hook generation.")
    add_visual(doc, "Ask the agent: 'what hook styles work best for me?'")
    add_narration(doc, "You can ask the agent 'what hook styles work best for me?' and it'll show you performance breakdowns by style. This data-driven approach takes the guesswork out of your video openings.")
    add_tip(doc, "Mention that hooks are stored per video, so you can compare performance across multiple videos over time.")

    # ============ VIDEO 21: MULTI-FORMAT PUBLISHING ============
    add_video_segment(doc, "21", "Multi-Format Publishing", "3 min",
        "How the system auto-generates 3 formats and publishes to all platforms with platform-specific captions.")

    add_visual(doc, "In the Library, show the format badges on a video: 16:9, 9:16, 1:1.")
    add_narration(doc, "Every video you upload to ContentForge is automatically generated in three formats: 16:9 horizontal for YouTube and Facebook, 9:16 vertical for TikTok and Instagram Reels, and 1:1 square for Instagram feed. You don't need to do anything — it's all automatic.")
    add_visual(doc, "Click 'Publish' on a ready video. Show the platform selection.")
    add_narration(doc, "When you publish, you select which platforms to post to. The system automatically uses the correct format for each platform — vertical for TikTok, horizontal for YouTube, and so on. No manual resizing needed.")
    add_visual(doc, "Explain the per-platform caption generation (mention this happens automatically).")
    add_narration(doc, "But it's not just the video format that's optimized. When you publish, AI generates platform-specific captions for each one. YouTube gets an SEO-optimized title and detailed description with 10 to 15 hashtags. TikTok gets a punchy title with emojis and 4 to 6 trending hashtags. Instagram gets an emoji-rich caption with 15 to 30 hashtags. Facebook gets conversational text. X gets a single tweet under 280 characters.")
    add_visual(doc, "Show the publish confirmation with 'platformCaptions' listed.")
    add_narration(doc, "After publishing, the confirmation shows which platforms received custom captions. Each platform gets content optimized for its culture and constraints — not the same copy pasted everywhere.")
    add_visual(doc, "Show the Scheduled tab with posts going to multiple platforms.")
    add_narration(doc, "If you schedule, each platform gets its own scheduled post at its own optimal time. So your TikTok might go out at noon, your YouTube at 3 PM, and your Instagram at 7 PM — all from a single video upload.")
    add_visual(doc, "End with a summary shot of the full workflow.")
    add_narration(doc, "And that's the ContentForge workflow: upload once, AI handles everything else — editing, captions, voiceover, multi-format generation, platform-specific copy, optimal scheduling, and publishing across all your social platforms. One video becomes optimized content for five platforms, on autopilot.")
    add_tip(doc, "This is a great closing video. End with a call to action — 'Start automating your content at ContentForge dot com' or similar.")

    # ============ CLOSING / OUTRO ============
    doc.add_page_break()
    add_heading(doc, "Outro Template (Use for All Videos)", level=1)
    add_para(doc, "Use this consistent outro for every video in the series:", italic=True)
    add_narration(doc, "And that's it for this feature! If you found this helpful, hit like and subscribe for more ContentForge tutorials. Check the description for links to other videos in this series, and visit ContentForge dot com to start automating your content today. See you in the next video!")
    add_para(doc, "")
    add_para(doc, "Standard outro elements to include:", bold=True)
    add_para(doc, "- Subscribe button overlay")
    add_para(doc, "- Link to next video in playlist")
    add_para(doc, "- Link to ContentForge platform")
    add_para(doc, "- Social media handles")
    add_para(doc, "- Background music fade-out")

    # ============ APPENDIX: AGENT COMMANDS ============
    doc.add_page_break()
    add_heading(doc, "Appendix: AI Agent Quick Commands", level=1)
    add_para(doc, "These are ready-to-use commands you can type into the AI Agent chat. Great for demo videos:", italic=True)

    commands = [
        ("What should I post this week?", "Generates a personalized weekly content strategy"),
        ("Give me 3 content ideas", "Generates 3 AI content ideas based on trends + analytics"),
        ("What's trending in my niche?", "Shows current trending hashtags and topics"),
        ("Schedule my ready videos at optimal times", "Auto-schedules all ready videos at platform-optimal times"),
        ("What videos do I have?", "Lists all videos in your library"),
        ("Show me my top performing videos", "Pulls analytics and ranks by views"),
        ("Why did my last video underperform?", "Triggers insights analysis and explains patterns"),
        ("Generate a thumbnail for my latest video", "Creates an AI thumbnail"),
        ("Repurpose my latest video into 5 clips", "Cuts a long video into short clips"),
        ("Translate my latest video to Spanish and French", "Translates captions to selected languages"),
        ("Generate 5 hooks for my latest video", "Creates 5 hook variants for A/B testing"),
        ("What are my competitors up to?", "Checks competitor posts for viral content"),
        ("Fetch new comments and suggest replies", "Pulls comments and generates AI replies"),
        ("What is my brand name and primary color?", "Shows brand kit settings"),
        ("Add @competitor_handle on YouTube as a competitor", "Adds a competitor to monitor"),
    ]

    for cmd, desc in commands:
        p = doc.add_paragraph()
        run = p.add_run(f"\"{cmd}\"")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xE8, 0x6A, 0x00)
        p.add_run(f" — {desc}").font.size = Pt(10)

    # ============ SAVE ============
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")
    print(f"Total videos: 21")
    print(f"Total runtime: ~75 minutes")

if __name__ == "__main__":
    build_document()
