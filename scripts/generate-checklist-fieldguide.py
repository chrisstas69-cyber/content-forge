#!/usr/bin/env python3
"""
ContentForge — Recording Checklist & Field Guide
A comprehensive manual covering pre-recording prep, recording setup,
post-recording steps, and a field-by-field explanation of every input.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUTPUT_PATH = "/home/z/my-project/download/ContentForge-Checklist-FieldGuide.docx"

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h

def add_para(doc, text, bold=False, italic=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_checkbox(doc, text, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("☐  ")
    run.font.size = Pt(12)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_field(doc, field_name, field_type, location, description, example=None, tips=None):
    """Add a field documentation block"""
    # Field name + type
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(field_name)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xE8, 0x6A, 0x00)
    run = p.add_run(f"  ({field_type})")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Location
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"Where: {location}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

    # Description
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"What it does: {description}")
    run.font.size = Pt(11)

    # Example
    if example:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"Example: ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(example)
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Tips
    if tips:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(f"💡 Tip: ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
        run = p.add_run(tips)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)

def build_document():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ============ COVER ============
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ContentForge")
    run.bold = True
    run.font.size = Pt(42)
    run.font.color.rgb = RGBColor(0xE8, 0x6A, 0x00)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Recording Checklist & Field Guide")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("Everything you need to record tutorial videos\nAND understand every field in the platform")
    run.font.size = Pt(13)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("Part 1: Recording Checklist (3 phases)\nPart 2: Field-by-Field Guide (18 tabs)")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ============ TABLE OF CONTENTS ============
    doc.add_page_break()
    add_heading(doc, "Table of Contents", level=1)

    add_para(doc, "PART 1: RECORDING CHECKLIST", bold=True, size=14, color=RGBColor(0xE8, 0x6A, 0x00))
    sections1 = [
        "Phase 1: Pre-Recording Setup (30 min)",
        "Phase 2: Recording Session (per video)",
        "Phase 3: Post-Recording (editing + publishing)",
        "Equipment & Software Recommendations",
    ]
    for s in sections1:
        add_para(doc, f"  • {s}", size=11)

    add_para(doc, "")
    add_para(doc, "PART 2: FIELD-BY-FIELD GUIDE", bold=True, size=14, color=RGBColor(0xE8, 0x6A, 0x00))
    sections2 = [
        "Dashboard Tab",
        "Upload Tab — Video Settings",
        "Upload Tab — Photo Slideshow Settings",
        "Upload Tab — AI Voiceover Settings",
        "Library Tab — Video Card",
        "Publish Dialog — Scheduling",
        "Ideas Tab",
        "Generate Tab — AI Thumbnails",
        "Generate Tab — AI B-roll Video",
        "Social Tab",
        "API Keys Tab",
        "Scheduled Tab",
        "Calendar Tab",
        "Trends Tab",
        "Analytics Tab",
        "Insights Tab",
        "Comments Tab",
        "Voice Tab",
        "Competitors Tab",
        "Brand Kit Tab",
        "Assets Tab",
        "Settings Tab",
        "AI Agent Chat",
    ]
    for s in sections2:
        add_para(doc, f"  • {s}", size=11)

    # ============ PART 1: RECORDING CHECKLIST ============
    doc.add_page_break()
    add_heading(doc, "PART 1: RECORDING CHECKLIST", level=1, color=RGBColor(0xE8, 0x6A, 0x00))

    # Phase 1
    add_heading(doc, "Phase 1: Pre-Recording Setup (30 minutes — do this ONCE)", level=2)
    add_para(doc, "Complete this checklist before recording your first video. You only need to do this once.", italic=True, size=10, color=RGBColor(0x66, 0x66, 0x66))

    add_para(doc, "1.1 Deploy ContentForge to Vercel", bold=True, size=12, space_after=4)
    add_checkbox(doc, "Push code to GitHub repository")
    add_checkbox(doc, "Import repo into Vercel (vercel.com/new)")
    add_checkbox(doc, "Set ENCRYPTION_KEY env var (run: openssl rand -hex 32)")
    add_checkbox(doc, "Set CRON_SECRET env var (run: openssl rand -hex 16)")
    add_checkbox(doc, "Switch database to Postgres (Vercel Storage → Create Postgres)")
    add_checkbox(doc, "Update prisma/schema.prisma: change 'sqlite' to 'postgresql'")
    add_checkbox(doc, "Redeploy and verify site loads at your Vercel URL")

    add_para(doc, "1.2 Set Up Your Content Environment", bold=True, size=12, space_after=4)
    add_checkbox(doc, "Open Settings tab → set Brand Handle (e.g. @yourdog)")
    add_checkbox(doc, "Open Settings tab → set Content Niche (e.g. 'dog / pet content')")
    add_checkbox(doc, "Open Brand Kit tab → upload your logo (PNG with transparency)")
    add_checkbox(doc, "Open Brand Kit tab → set primary color, font, watermark position")
    add_checkbox(doc, "Open Voice tab → fill in Persona, Tone, Signature Phrases")
    add_checkbox(doc, "Open Assets tab → upload watermark, intro clip, outro clip, background music")
    add_checkbox(doc, "Open API Keys tab → add at least YouTube credentials (for demos)")
    add_checkbox(doc, "Open Social tab → connect at least one platform via OAuth")
    add_checkbox(doc, "Upload 2-3 real videos (so Library has content to show)")
    add_checkbox(doc, "Upload 1-2 photos as a slideshow (to demo photo-to-video)")
    add_checkbox(doc, "Wait for all videos to reach 'ready' status")
    add_checkbox(doc, "Go to Ideas tab → click 'Generate 5 ideas'")
    add_checkbox(doc, "Go to Trends tab → click 'Refresh trends'")
    add_checkbox(doc, "Go to Competitors tab → add 1-2 competitors and click 'Check now'")

    add_para(doc, "1.3 Prepare Your Recording Environment", bold=True, size=12, space_after=4)
    add_checkbox(doc, "Close all unnecessary browser tabs and applications")
    add_checkbox(doc, "Hide browser bookmarks bar (Ctrl+Shift+B in Chrome)")
    add_checkbox(doc, "Set browser to full screen / maximize window")
    add_checkbox(doc, "Set screen resolution to 1920x1080 or higher")
    add_checkbox(doc, "Enable screen magnifier tool (for zooming into UI elements)")
    add_checkbox(doc, "Turn off notifications (Do Not Disturb mode)")
    add_checkbox(doc, "Test microphone — record 10 seconds and playback")
    add_checkbox(doc, "Position microphone 6-8 inches from mouth")
    add_checkbox(doc, "Have the video script document open on a second monitor (or printed)")
    add_checkbox(doc, "Have a glass of water nearby")
    add_checkbox(doc, "Set browser zoom to 100% (Ctrl+0)")
    add_checkbox(doc, "Use a clean browser profile (no extensions visible except essentials)")

    add_para(doc, "1.4 Test Recording Software", bold=True, size=12, space_after=4)
    add_checkbox(doc, "Install OBS Studio (free) or your preferred screen recorder")
    add_checkbox(doc, "Set recording quality to 1080p, 30fps (or 4K if your screen supports it)")
    add_checkbox(doc, "Set audio sample rate to 48kHz")
    add_checkbox(doc, "Do a 30-second test recording — verify video + audio quality")
    add_checkbox(doc, "Check file format (MP4 or MKV — MP4 is easier to edit)")

    # Phase 2
    doc.add_page_break()
    add_heading(doc, "Phase 2: Recording Session (per video)", level=2)
    add_para(doc, "Follow this checklist for EACH video you record.", italic=True, size=10, color=RGBColor(0x66, 0x66, 0x66))

    add_para(doc, "2.1 Before Hitting Record", bold=True, size=12, space_after=4)
    add_checkbox(doc, "Read the video script segment from the script document")
    add_checkbox(doc, "Navigate to the starting tab/section in ContentForge")
    add_checkbox(doc, "Prepare any data you'll need (e.g. have a video ready to upload)")
    add_checkbox(doc, "Clear any test/placeholder data that shouldn't be in the video")
    add_checkbox(doc, "Take a sip of water")
    add_checkbox(doc, "Deep breath — relax your shoulders")
    add_checkbox(doc, "Start screen recorder")
    add_checkbox(doc, "Wait 3 seconds before speaking (gives clean intro for editing)")

    add_para(doc, "2.2 During Recording", bold=True, size=12, space_after=4)
    add_checkbox(doc, "Speak clearly and at a moderate pace (not too fast)")
    add_checkbox(doc, "Move mouse slowly when pointing at UI elements (viewers need to follow)")
    add_checkbox(doc, "Pause 2 seconds after clicking buttons (let the UI respond)")
    add_checkbox(doc, "If you make a mistake, pause 3 seconds, then re-do the segment (easy to edit)")
    add_checkbox(doc, "Zoom into important UI elements using screen magnifier")
    add_checkbox(doc, "Don't read the script word-for-word — paraphrase naturally")
    add_checkbox(doc, "End with the standard outro ('hit like and subscribe…')")
    add_checkbox(doc, "Wait 3 seconds after finishing before stopping recording")

    add_para(doc, "2.3 After Each Video", bold=True, size=12, space_after=4)
    add_checkbox(doc, "Stop recording")
    add_checkbox(doc, "Save file with descriptive name (e.g. '03-uploading-videos-take1.mp4')")
    add_checkbox(doc, "Watch the first 30 seconds to verify audio + video quality")
    add_checkbox(doc, "Note any segments that need re-recording")
    add_checkbox(doc, "Reset any demo data if needed (delete test uploads, etc.)")

    # Phase 3
    doc.add_page_break()
    add_heading(doc, "Phase 3: Post-Recording (editing + publishing)", level=2)

    add_para(doc, "3.1 Editing", bold=True, size=12, space_after=4)
    add_checkbox(doc, "Import raw recordings into video editor (CapCut, DaVinci Resolve, Premiere)")
    add_checkbox(doc, "Cut out mistakes and long pauses")
    add_checkbox(doc, "Add zoom effects on key UI elements (use keyframe zoom)")
    add_checkbox(doc, "Add intro/outro graphics (consistent across all videos)")
    add_checkbox(doc, "Add background music (low volume, -20dB below narration)")
    add_checkbox(doc, "Add captions/subtitles (use auto-caption tool, then fix errors)")
    add_checkbox(doc, "Color-correct if needed (brightness/contrast)")
    add_checkbox(doc, "Export at 1080p MP4, H.264, 10-15 Mbps bitrate")

    add_para(doc, "3.2 Publishing", bold=True, size=12, space_after=4)
    add_checkbox(doc, "Create YouTube playlist: 'ContentForge Tutorial Series'")
    add_checkbox(doc, "Upload video with title: 'ContentForge: [Feature Name] Tutorial'")
    add_checkbox(doc, "Write description with timestamps + link to platform")
    add_checkbox(doc, "Add thumbnail (use the AI thumbnail generator!)")
    add_checkbox(doc, "Add to playlist")
    add_checkbox(doc, "Set visibility: Public or Unlisted (your choice)")
    add_checkbox(doc, "Add end screen pointing to next video in series")

    add_para(doc, "3.3 Repurposing Tutorial Content", bold=True, size=12, space_after=4)
    add_checkbox(doc, "Cut 30-60 second highlights for TikTok/Reels")
    add_checkbox(doc, "Write a Twitter thread summarizing key points")
    add_checkbox(doc, "Create a blog post version (transcribe + format)")
    add_checkbox(doc, "Post a LinkedIn carousel with screenshots")

    # Equipment recommendations
    doc.add_page_break()
    add_heading(doc, "Equipment & Software Recommendations", level=2)

    add_para(doc, "Screen Recording Software", bold=True, size=12, space_after=4)
    add_para(doc, "• OBS Studio (FREE) — Best overall, works on Windows/Mac/Linux. Bit of a learning curve but extremely powerful.")
    add_para(doc, "• Loom (FREE tier) — Easiest to use, browser-based. Good for quick recordings. Limited to 5 min on free plan.")
    add_para(doc, "• QuickTime Player (FREE, Mac only) — Simple, built into macOS. File → New Screen Recording.")
    add_para(doc, "• Camtasia ($299) — Professional screen recording + editing in one. Best if you record tutorials frequently.")

    add_para(doc, "Microphone", bold=True, size=12, space_after=4)
    add_para(doc, "• Blue Yeti USB Mic (~$130) — Industry standard for YouTubers. Plug and play.")
    add_para(doc, "• Rode NT-USB (~$170) — Higher quality, still USB-simple.")
    add_para(doc, "• Your AirPods/headphones mic — Okay for starting out, not ideal for professional quality.")
    add_para(doc, "• Built-in laptop mic — Last resort. Audio quality matters more than video quality.")

    add_para(doc, "Video Editing Software", bold=True, size=12, space_after=4)
    add_para(doc, "• CapCut (FREE) — Best free editor. Works on desktop + mobile. Great for social-first content.")
    add_para(doc, "• DaVinci Resolve (FREE) — Professional-grade editor. Steep learning curve but extremely powerful.")
    add_para(doc, "• iMovie (FREE, Mac only) — Simple, good for beginners.")
    add_para(doc, "• Adobe Premiere Pro ($22/mo) — Industry standard if you already have Creative Cloud.")

    add_para(doc, "Helpful Tools", bold=True, size=12, space_after=4)
    add_para(doc, "• Screen Magnifier (built into Windows/Mac) — For zooming into UI elements during recording")
    add_para(doc, "• Grammarly (FREE) — For proofreading video descriptions")
    add_para(doc, "• Canva (FREE) — For creating custom intro/outro graphics")
    add_para(doc, "• TubeBuddy or VidIQ (FREE) — For YouTube SEO optimization")

    # ============ PART 2: FIELD-BY-FIELD GUIDE ============
    doc.add_page_break()
    add_heading(doc, "PART 2: FIELD-BY-FIELD GUIDE", level=1, color=RGBColor(0xE8, 0x6A, 0x00))
    add_para(doc, "This section explains every input field, dropdown, toggle, and button in ContentForge. Use it as your reference manual while learning the platform and recording tutorials.", italic=True, size=11)

    # ---- DASHBOARD ----
    doc.add_page_break()
    add_heading(doc, "Dashboard Tab", level=2)
    add_para(doc, "The Dashboard is your overview screen. No inputs here — just displays. But understand what each stat means.", italic=True, size=10)

    add_field(doc, "Total Videos", "Stat card", "Dashboard → top row",
        "Total count of all videos in your library, regardless of status.",
        "If it shows '5', you have 5 videos total.")
    add_field(doc, "Processing", "Stat card", "Dashboard → top row",
        "Number of videos currently being edited by FFmpeg or analyzed by AI.",
        "Shows '2' means 2 videos are in the editing pipeline.")
    add_field(doc, "Ready", "Stat card", "Dashboard → top row",
        "Videos that have completed processing and are ready to publish.",
        "Shows '3' means 3 videos are ready to go.")
    add_field(doc, "Scheduled", "Stat card", "Dashboard → top row",
        "Posts that have been scheduled but haven't published yet.",
        "Shows '5' means 5 posts are queued for future publishing.")
    add_field(doc, "Published", "Stat card", "Dashboard → top row",
        "Videos that have been successfully published to at least one platform.",
        "Shows '10' means 10 videos have gone live.")
    add_field(doc, "Avg Viral Score", "Stat card", "Dashboard → top row",
        "Average viral score (0-100) across all your scored videos. Updates in real-time.",
        "Shows '78' means your average viral potential is 78/100.")
    add_field(doc, "Quick Actions", "Buttons", "Dashboard → middle section",
        "Shortcuts to common tasks. Click any button to navigate to that feature.",
        "Click 'Upload new videos' to jump to the Upload tab.")
    add_field(doc, "Connected Accounts", "Display", "Dashboard → right panel",
        "Shows how many social accounts are connected and on which platforms.",
        "Shows 'YouTube: 1 connected' means your YouTube account is linked.")

    # ---- UPLOAD TAB ----
    doc.add_page_break()
    add_heading(doc, "Upload Tab — Video Settings", level=2)
    add_para(doc, "Fields for uploading video files and configuring how they're processed.", italic=True, size=10)

    add_field(doc, "Video Drop Zone", "File input", "Upload tab → top",
        "Click or drag video files here. Accepts MP4, MOV, WebM. Multiple files allowed.",
        "Drag 3 MP4 files here to upload all 3 at once.")
    add_field(doc, "File List", "Display + remove", "Upload tab → below dropzone",
        "Shows selected files before upload. Click X to remove any file.",
        "Shows 'video1.mp4, video2.mp4' with X buttons next to each.")
    add_field(doc, "Burn AI Captions", "Toggle", "Upload tab → Auto-Editing Options",
        "When ON (default), AI transcribes your video's audio and burns subtitles onto the video. Great for accessibility and engagement.",
        "ON = subtitles appear on the final video.",
        "Keep ON for most content. Turn OFF only if you'll add captions in post-production.")
    add_field(doc, "Auto-trim Silence", "Toggle", "Upload tab → Auto-Editing Options",
        "When ON, FFmpeg detects and removes silent segments from the video. Reduces dead air.",
        "ON = a 60s video with 10s of silence becomes 50s.",
        "Good for talking-head videos. Turn OFF for music/atmospheric content where silence is intentional.")
    add_field(doc, "Watermark", "Dropdown", "Upload tab → Auto-Editing Options",
        "Select a previously-uploaded PNG logo to overlay on the video. Upload logos in the Assets tab first.",
        "Select 'MyLogo.png' from the dropdown to add it to every frame.",
        "PNG with transparency works best. 300x300px is a good size.")
    add_field(doc, "Watermark Position", "Dropdown", "Upload tab → Auto-Editing Options",
        "Where on the video the watermark appears. Options: Top Left, Top Right, Bottom Left, Bottom Right, Center.",
        "Select 'Bottom Right' to put your logo in the corner.",
        "Bottom Right is the standard for most platforms. Avoid Center — it blocks content.")
    add_field(doc, "Background Music", "Dropdown", "Upload tab → Auto-Editing Options",
        "Select a previously-uploaded music file to mix under your video. Upload music in the Assets tab first.",
        "Select 'upbeat_track.mp3' to add background music.",
        "Music plays at 20% volume by default. Royalty-free music only — don't use copyrighted songs.")
    add_field(doc, "Intro Clip", "Dropdown", "Upload tab → Auto-Editing Options",
        "Select a video clip to prepend to the start of your video. Upload intros in the Assets tab first.",
        "Select 'intro_3sec.mp4' to add a 3-second branded intro.",
        "Keep intros under 5 seconds. Longer intros cause viewers to scroll past.")
    add_field(doc, "Outro Clip", "Dropdown", "Upload tab → Auto-Editing Options",
        "Select a video clip to append to the end of your video. Upload outros in the Assets tab first.",
        "Select 'outro_5sec.mp4' to add a call-to-action outro.",
        "Use outros for 'subscribe' or 'follow for more' CTAs.")

    # ---- PHOTO SETTINGS ----
    doc.add_page_break()
    add_heading(doc, "Upload Tab — Photo Slideshow Settings", level=2)
    add_para(doc, "These fields appear when you upload photos instead of videos.", italic=True, size=10)

    add_field(doc, "Upload Photos Instead", "Button", "Upload tab → below video dropzone",
        "Click to open a file picker for images. Accepts PNG, JPG. Multiple files allowed.",
        "Click → select 4 photos → they appear in the purple 'Photos → Slideshow' section.")
    add_field(doc, "Seconds per Photo", "Slider (2-10s)", "Upload tab → Photo Slideshow Settings",
        "How long each photo displays before transitioning to the next. Default: 5 seconds.",
        "Set to 3s for fast-paced slideshows, 7s for slower, dramatic pacing.",
        "For TikTok/Reels, keep it fast (3-4s). For YouTube, 5-7s is fine.")
    add_field(doc, "Transition", "Slider (0.3-2s)", "Upload tab → Photo Slideshow Settings",
        "Duration of the crossfade transition between photos. Default: 0.7 seconds.",
        "0.3s = quick cut, 1.5s = slow dreamy fade.",
        "Shorter transitions feel more energetic. Longer feels more cinematic.")
    add_field(doc, "Custom Voiceover Script", "Textarea", "Upload tab → Photo Slideshow Settings",
        "Optional — write your own script for the AI voiceover. If blank, AI generates one based on your niche + trends.",
        "Type 'Meet Max, the goodest boy! Every morning he wakes me up at 6am...' to use that exact text.",
        "If you provide a script, it's used for BOTH the voiceover audio AND the burned captions. Maximum control.")

    # ---- VOICEOVER SETTINGS ----
    doc.add_page_break()
    add_heading(doc, "Upload Tab — AI Voiceover Settings", level=2)

    add_field(doc, "Enable AI Voiceover", "Toggle", "Upload tab → AI Voiceover section",
        "When ON, AI generates a voiceover script and synthesizes it as audio. Mixed over the video.",
        "ON = AI writes a script + speaks it. OFF = no voiceover.",
        "Auto-enabled for photo uploads. For videos, turn ON if you want narration.")
    add_field(doc, "Voice", "Dropdown", "Upload tab → AI Voiceover section",
        "Select from 7 AI voices. Each has a different gender and tone.",
        "Tongtong = warm neutral. Tianmei = cheerful female. Yunlong = deep male.",
        "Try different voices to find one that matches your brand. Test with a short video first.")
    add_field(doc, "Tone", "Text input", "Upload tab → AI Voiceover section",
        "Describe the tone you want the script to have. AI uses this when writing the voiceover.",
        "'funny, energetic, engaging' or 'calm, professional, informative'",
        "Be specific. 'Funny' alone is vague. 'Funny with dad jokes and dog puns' is better.")
    add_field(doc, "Replace Original Audio", "Toggle", "Upload tab → AI Voiceover section",
        "When ON, voiceover replaces all original audio. When OFF, voiceover mixes on top of existing audio.",
        "ON = only voiceover is heard. OFF = voiceover + original audio.",
        "For photo uploads, always ON (no original audio). For videos, OFF usually sounds better.")

    # ---- LIBRARY ----
    doc.add_page_break()
    add_heading(doc, "Library Tab — Video Card", level=2)
    add_para(doc, "Understanding the elements on each video card in your Library.", italic=True, size=10)

    add_field(doc, "Viral Score Circle", "Display (color-coded)", "Library → video card → top right",
        "AI-generated score 0-100. Green = 80+ (high viral potential). Amber = 50-79. Red = below 50.",
        "Shows '85' in a green circle = high viral potential.",
        "The score is based on hook strength, emotional pull, trend alignment, and rewatchability.")
    add_field(doc, "Status Badge", "Display", "Library → video card → top left",
        "Shows current pipeline status: pending, editing, transcribing, scoring, ready, published, failed.",
        "Shows 'ready' in green = video is done and ready to publish.",
        "If 'failed', click Retry to reprocess. Check errorMessage for details.")
    add_field(doc, "Clip Badge", "Display", "Library → video card → top left",
        "Purple badge indicating this video is a clip extracted from a longer video via Repurpose.",
        "Shows 'Clip' in purple = this is a short clip from a parent video.",
        "Clip cards also show 'from 45s' meaning it starts at the 45-second mark of the original.")
    add_field(doc, "Format Badges", "Display", "Library → video card → below preview",
        "Green badges showing which aspect ratios have been generated: 16:9, 9:16, 1:1.",
        "Shows '✓ 16:9, ✓ 9:16, ✓ 1:1' = all three formats are ready.",
        "If badges are missing, the video is still processing formats.")
    add_field(doc, "Publish Button", "Button", "Library → video card → bottom",
        "Opens the publish dialog where you select platforms, edit captions, and schedule. Only appears when video is 'ready'.",
        "Click Publish → dialog opens with platform selection + scheduling options.")
    add_field(doc, "Repurpose Button", "Button", "Library → video card → bottom",
        "Purple button that triggers AI to cut a long video into 5 short clips. Only appears on videos longer than 90 seconds.",
        "Click Repurpose → AI identifies engaging moments → extracts 5 clips.",
        "Each clip becomes its own video in the Library with a purple 'Clip' badge.")
    add_field(doc, "Translate Button", "Button", "Library → video card → bottom",
        "Opens the translate dialog where you select languages for caption translation. Supports 12 languages.",
        "Click Translate → select Spanish, French, Japanese → click 'Translate to 3'.")
    add_field(doc, "Retry Button", "Button", "Library → video card → bottom",
        "Appears when video status is 'failed'. Re-triggers the processing pipeline from the beginning.",
        "Click Retry → video goes back to 'pending' and reprocesses.")
    add_field(doc, "Delete Button", "Button", "Library → video card → bottom",
        "Permanently deletes the video and all associated files. Asks for confirmation.",
        "Click Delete → 'Delete this video?' → click OK.")

    # ---- PUBLISH DIALOG ----
    doc.add_page_break()
    add_heading(doc, "Publish Dialog — Scheduling", level=2)

    add_field(doc, "Title", "Text input", "Publish dialog",
        "The title for your post. Pre-filled with AI-generated title. Edit as needed.",
        "AI suggests 'When your dog thinks they're the boss' — you can change it.",
        "Each platform gets its own optimized title variant automatically.")
    add_field(doc, "Description", "Textarea", "Publish dialog",
        "Post description. Pre-filled with AI-generated description. Edit as needed.",
        "AI suggests a 2-3 sentence description. You can expand or rewrite.")
    add_field(doc, "Hashtags", "Text input (comma-separated)", "Publish dialog",
        "Hashtags for your post. Pre-filled with AI suggestions. Add or remove as needed.",
        "'dogsoftiktok, puppylove, funnydogs, dogmom'",
        "Each platform gets its own optimized hashtag set (TikTok gets fewer, IG gets more).")
    add_field(doc, "Publish to", "Multi-select buttons", "Publish dialog",
        "Select which platforms to publish to. Shows only connected accounts.",
        "Click YouTube + TikTok → both get selected (green border).")
    add_field(doc, "When to publish", "3-button toggle", "Publish dialog → bottom",
        "Choose: 'Publish now' (immediate), 'Optimal times' (AI picks best time per platform), or 'Pick time' (manual datetime).",
        "Select 'Optimal times' → each platform gets scheduled at its peak engagement hour.",
        "Optimal times uses industry data. Once you have analytics, it uses YOUR data instead.")
    add_field(doc, "Datetime Picker", "Datetime input", "Publish dialog → when 'Pick time' selected",
        "Choose a specific date and time to publish. Appears only when 'Pick time' is selected.",
        "Select '2026-07-01 14:00' to publish on July 1st at 2pm.")

    # ---- IDEAS ----
    doc.add_page_break()
    add_heading(doc, "Ideas Tab", level=2)

    add_field(doc, "Generate 5 ideas", "Button", "Ideas tab → top right",
        "Triggers AI to generate 5 content ideas based on trends, analytics, and your niche. Takes 15-20 seconds.",
        "Click → wait 20s → 5 idea cards appear.",
        "Generates different ideas each time. Click multiple times for more options.")
    add_field(doc, "Idea Card — Title", "Display", "Ideas tab → idea card",
        "AI-generated scroll-stopping title for a content idea.",
        "'Paw-sonality Test: What Dog Are You?'")
    add_field(doc, "Idea Card — Predicted Viral Score", "Display", "Ideas tab → idea card",
        "AI's prediction of how viral this idea could be (0-100). Color-coded.",
        "Shows '92' in large text = very high potential.")
    add_field(doc, "Idea Card — Format", "Display", "Ideas tab → idea card",
        "Recommended aspect ratio: 9:16 (TikTok/Reels), 16:9 (YouTube), or 1:1 (IG feed).",
        "Shows '9:16' = make this as a vertical video.")
    add_field(doc, "Idea Card — Source", "Display", "Ideas tab → idea card",
        "What data informed this idea: trends, insights, analytics, or manual.",
        "Shows 'Trends' = based on current trending content.")
    add_field(doc, "Idea Card — Script Outline", "Display (expandable)", "Ideas tab → idea card",
        "Hook + Body + CTA breakdown for the idea. Production-ready outline.",
        "Hook: 'Quick question - if you were a dog, which breed would you be?'")
    add_field(doc, "Dismiss", "Button", "Ideas tab → idea card → bottom",
        "Removes the idea from your 'Fresh ideas' list. Idea is deleted.",
        "Click Dismiss → idea disappears from the list.")

    # ---- GENERATE TAB ----
    doc.add_page_break()
    add_heading(doc, "Generate Tab — AI Thumbnails", level=2)

    add_field(doc, "AI Thumbnail / AI Image / AI B-roll Video", "Tab selector", "Generate tab → top",
        "Switch between 3 generation modes. Click any card to select.",
        "Click 'AI Thumbnail' card → thumbnail generation mode is active.")
    add_field(doc, "Video title", "Text input", "Generate tab → AI Thumbnail mode",
        "The video title to base the thumbnail on. AI uses this as inspiration.",
        "Type 'Puppy's First Day at the Beach' → AI generates a beach-themed thumbnail.")
    add_field(doc, "Target video", "Dropdown", "Generate tab → AI Thumbnail mode",
        "Optional — link the thumbnail to a specific video in your library.",
        "Select 'Puppy's First Snow' from dropdown → thumbnail is associated with that video.")
    add_field(doc, "Upload your own image", "File input", "Generate tab → AI Thumbnail mode",
        "Optional — upload a photo (e.g. your dog) and AI transforms it into a stylized thumbnail using Replicate SDXL img2img.",
        "Upload 'my_dog.jpg' → AI creates a thumbnail using your dog's image.",
        "Requires Replicate API token in API Keys tab. Without upload, uses built-in AI (no token needed).")
    add_field(doc, "Style Strength", "Slider (10-80%)", "Generate tab → AI Thumbnail mode (when image uploaded)",
        "Controls how much AI changes your uploaded photo. Lower = more recognizable. Higher = more AI transformation.",
        "35% (default) = balanced mix of your photo + thumbnail styling.",
        "Start at 35%. Try 20% for subtle styling, 50% for dramatic transformation.")
    add_field(doc, "Generate image / Generate from your photo", "Button", "Generate tab → bottom",
        "Triggers generation. Text changes based on whether you uploaded an image.",
        "Click 'Generate from your photo' → processing starts (30-60 seconds).")

    add_heading(doc, "Generate Tab — AI B-roll Video", level=2)
    add_field(doc, "Prompt", "Textarea", "Generate tab → AI B-roll mode",
        "Describe the video clip you want AI to generate.",
        "'cinematic shot of a golden retriever running on a beach at sunset, slow motion'",
        "Be descriptive. Include camera style, subject, setting, and mood. Requires Replicate API token.")
    add_field(doc, "Generate video", "Button", "Generate tab → AI B-roll mode",
        "Triggers video generation. Takes 2-5 minutes. Runs in background.",
        "Click → 'Video generation started' message appears.")

    # ---- SOCIAL ----
    doc.add_page_break()
    add_heading(doc, "Social Tab", level=2)

    add_field(doc, "Platform Status Badge", "Display", "Social tab → each platform card",
        "Shows 'API Ready' (green) if credentials are set, or 'Needs Keys' (amber) if not.",
        "Shows 'API Ready' → you can connect this platform.",
        "If 'Needs Keys', click 'Add Keys' to go to the API Keys tab.")
    add_field(doc, "Connect", "Button", "Social tab → platform card",
        "Starts the OAuth flow. Redirects you to the platform's login page. Only appears when API is ready.",
        "Click Connect → redirected to Google → authorize → redirected back.",
        "Make sure the OAuth Redirect URL matches exactly what's in your developer console.")
    add_field(doc, "Disconnect", "Button", "Social tab → platform card",
        "Marks the account as disconnected. You'll need to reconnect to publish again.",
        "Click Disconnect → account shows 'Not connected'.")

    # ---- API KEYS ----
    doc.add_page_break()
    add_heading(doc, "API Keys Tab", level=2)

    add_field(doc, "Platform Group Card", "Display", "API Keys tab → per platform",
        "Shows platform name, description, 'Get keys' link, and 'Configured'/'Not configured' badge.",
        "YouTube card shows 'Not configured' → click 'Get keys' to open Google Cloud Console.")
    add_field(doc, "Get keys", "Link", "API Keys tab → platform card",
        "Opens the platform's developer portal in a new tab where you can create an app and get credentials.",
        "Click 'Get keys' next to YouTube → opens console.cloud.google.com.")
    add_field(doc, "Client ID / Client Secret / API Token", "Password input", "API Keys tab → platform card",
        "Paste your API credentials here. Values are encrypted before storage. Shows masked preview after saving.",
        "Paste '1234567890-abc.apps.googleusercontent.com' → saved as '1234••••••••••.com'",
        "Never share these values. The masked preview is the most anyone will ever see.")
    add_field(doc, "Save [Platform] keys", "Button", "API Keys tab → platform card",
        "Saves the entered credentials. Button is disabled until at least one field has a value.",
        "Click 'Save YouTube keys' → credentials encrypted + stored → 'Configured' badge appears.")
    add_field(doc, "Trash icon", "Button", "API Keys tab → saved credential",
        "Deletes a saved credential. You'll need to re-enter it to use that platform.",
        "Click trash icon → 'Remove Client ID?' → click OK.")

    # ---- SCHEDULED ----
    doc.add_page_break()
    add_heading(doc, "Scheduled Tab", level=2)

    add_field(doc, "Summary Cards", "Display", "Scheduled tab → top",
        "Four cards showing counts: Scheduled, In progress, Published (24h), Failed (24h).",
        "Shows 'Scheduled: 5' = 5 posts are queued for future publishing.")
    add_field(doc, "Upcoming / Publishing now / Recently failed / Recently published", "Display sections", "Scheduled tab",
        "Posts grouped by status. Each shows thumbnail, title, platform, scheduled time, and status badge.",
        "Upcoming section shows: 'Puppy Video → TikTok → Fri Jul 5, 7:00 PM'.")
    add_field(doc, "Cancel", "Button", "Scheduled tab → scheduled post card",
        "Cancels a scheduled post before it publishes. Only appears on 'scheduled' posts.",
        "Click Cancel → post is removed from the queue.")

    # ---- CALENDAR ----
    doc.add_page_break()
    add_heading(doc, "Calendar Tab", level=2)

    add_field(doc, "Month Navigation", "Buttons (← →)", "Calendar tab → top right",
        "Navigate between months. Shows current month name and year.",
        "Click → → moves to next month.")
    add_field(doc, "Day Cells", "Display grid", "Calendar tab → main area",
        "Calendar grid showing each day of the month. Today is highlighted in orange. Posts appear as colored dots.",
        "July 15 shows blue dot 'TikTok · 7:00 PM' = scheduled post on that day.")
    add_field(doc, "Post Dots", "Display (color-coded)", "Calendar tab → day cells",
        "Colored dots showing posts per day. Blue = scheduled, green = published, red = failed, amber = uploading.",
        "Blue dot = post is scheduled but not yet published.")

    # ---- TRENDS ----
    doc.add_page_break()
    add_heading(doc, "Trends Tab", level=2)

    add_field(doc, "Refresh trends", "Button", "Trends tab → top right",
        "Triggers a web search for trending content in your niche. Takes 30-60 seconds.",
        "Click → 'Searching web…' → results populate after ~45 seconds.",
        "Trends also auto-refresh daily via cron. Manual refresh gets the latest data.")
    add_field(doc, "Platform Filter", "Buttons", "Trends tab → below header",
        "Filter trends by platform: All platforms, TikTok, Instagram, YouTube, X.",
        "Click 'TikTok (10)' → shows only TikTok trends.",
        "Number in parentheses shows how many trends exist for that platform.")
    add_field(doc, "Trend Card", "Display", "Trends tab → list",
        "Shows trend type (hashtag/sound/format/topic), content, summary, score (0-100), and discovery date.",
        "Shows '#dogsoftiktok · HASHTAG · Score: 95 · 6/25/2026'.")
    add_field(doc, "Score", "Display (color-coded)", "Trends tab → trend card",
        "Relevance score 0-100. Higher = more relevant to your niche.",
        "Score 95 (green icon) = very relevant. Score 65 (amber) = moderately relevant.")

    # ---- ANALYTICS ----
    doc.add_page_break()
    add_heading(doc, "Analytics Tab", level=2)

    add_field(doc, "Total Views / Likes / Comments / Shares", "Stat cards", "Analytics tab → top",
        "Aggregate metrics across all platforms and posts.",
        "Shows '12,345' views = total views across all your published posts.")
    add_field(doc, "Refresh metrics", "Button", "Analytics tab → top right",
        "Pulls fresh metrics from all connected platforms. Takes 10-30 seconds.",
        "Click → 'Refreshing…' → updated metrics appear.",
        "Also auto-refreshes hourly via cron.")
    add_field(doc, "Per-Platform Breakdown", "Display", "Analytics tab → middle",
        "Shows post count and views per platform.",
        "YouTube: 5 posts, 8,000 views. TikTok: 3 posts, 4,000 views.")
    add_field(doc, "Top Performing Videos", "Display (ranked)", "Analytics tab → bottom",
        "Your top 10 videos ranked by total views. Shows rank, thumbnail, title, and metrics.",
        "#1: 'Puppy's First Snow' — 5,000 views, 500 likes.")

    # ---- INSIGHTS ----
    doc.add_page_break()
    add_heading(doc, "Insights Tab", level=2)

    add_field(doc, "Refresh insights", "Button", "Insights tab → top right",
        "Triggers AI to analyze your last 30 videos + analytics and generate insights. Takes ~20 seconds.",
        "Click → 'Analyzing…' → 5-8 insights appear.",
        "Also auto-refreshes daily via cron. Old insights are replaced with fresh ones.")
    add_field(doc, "Insight Card", "Display (color-coded)", "Insights tab → list",
        "Shows insight type (Pattern/Opportunity/Underperformer/Recommendation), content, confidence %, and supporting data.",
        "Pattern (green): 'Your puppy videos get 3x more views than training videos.'")
    add_field(doc, "Confidence", "Display", "Insights tab → insight card",
        "How certain the AI is about this insight (0-100%). Higher = more data supports it.",
        "Confidence: 85% = strong pattern backed by multiple data points.")
    add_field(doc, "view data", "Expandable", "Insights tab → insight card",
        "Click to see the supporting metrics behind the insight.",
        "Click 'view data' → shows JSON with avgPuppyViews, avgTrainingViews, etc.")

    # ---- COMMENTS ----
    doc.add_page_break()
    add_heading(doc, "Comments Tab", level=2)

    add_field(doc, "Fetch new comments", "Button", "Comments tab → top right",
        "Pulls recent comments from all published posts and generates AI replies. Takes ~30 seconds.",
        "Click → 'Fetching…' → 'Fetched 12 comments, 12 replies suggested'.",
        "Also auto-runs 4x daily via cron.")
    add_field(doc, "Filter (Pending/Replied/Ignored/All)", "Buttons", "Comments tab → below header",
        "Filter comments by reply status.",
        "Click 'Pending' → shows only comments awaiting your review.")
    add_field(doc, "Comment Card", "Display", "Comments tab → list",
        "Shows author, platform, comment text, video title, video thumbnail, and AI-suggested reply in blue box.",
        "Shows 'John Doe · YouTube · So cute!' with AI reply 'Thanks! He really is 😊'.")
    add_field(doc, "Approve & Post", "Button", "Comments tab → comment card",
        "Posts the AI-suggested reply immediately to the platform.",
        "Click → reply is published → status changes to 'Replied'.")
    add_field(doc, "Edit", "Button", "Comments tab → comment card",
        "Opens the suggested reply for editing before posting.",
        "Click → textarea appears with suggested text → modify → click 'Post'.")
    add_field(doc, "Ignore", "Button", "Comments tab → comment card",
        "Marks the comment as ignored. No reply is posted.",
        "Click → comment moves to 'Ignored' filter.")

    # ---- VOICE ----
    doc.add_page_break()
    add_heading(doc, "Voice Tab", level=2)

    add_field(doc, "Persona", "Text input", "Voice tab",
        "Describe yourself in a sentence. AI uses this to match your voice in replies and scripts.",
        "'friendly dog mom who loves puns' or 'energetic fitness coach'.",
        "Be specific. The more detail you give, the better AI matches your style.")
    add_field(doc, "Tone", "Text input", "Voice tab",
        "Describe the emotional tone you want. Controls how replies 'feel'.",
        "'warm, witty, enthusiastic' or 'calm, professional, informative'.")
    add_field(doc, "Signature Phrases", "Text input (comma-separated)", "Voice tab",
        "Words or expressions you use frequently. AI weaves these into replies naturally.",
        "'pawsome, fur-real, zoomies' or 'let's go, no excuses, grind time'.",
        "Don't overdo it. 3-5 phrases is plenty. Too many sounds forced.")
    add_field(doc, "Phrases to Avoid", "Text input", "Voice tab",
        "Words you never want AI to use. Blacklist.",
        "'literally, basically, obviously'.")
    add_field(doc, "Reply Mode", "Dropdown", "Voice tab",
        "'Suggest' = generate replies for manual approval (recommended). 'Auto' = post replies immediately without approval.",
        "Select 'Suggest' → you review each reply before it posts.",
        "Start with 'Suggest'. Switch to 'Auto' only when you trust the AI and have high comment volume.")
    add_field(doc, "Reply Length", "Dropdown", "Voice tab",
        "How long AI replies are: 'Short' (1 sentence), 'Medium' (2-3 sentences), 'Long' (3-5 sentences).",
        "Select 'Short' → replies are punchy one-liners.",
        "Most platforms favor short replies. Use 'Medium' for thoughtful responses to questions.")
    add_field(doc, "Save Voice Profile", "Button", "Voice tab → bottom",
        "Saves your voice profile. Immediately affects all AI-generated content.",
        "Click → 'Voice profile saved' toast appears.")

    # ---- COMPETITORS ----
    doc.add_page_break()
    add_heading(doc, "Competitors Tab", level=2)

    add_field(doc, "Platform Dropdown", "Dropdown", "Competitors tab → add form",
        "Select which platform the competitor is on: YouTube, Instagram, TikTok, or X.",
        "Select 'youtube' → competitor will be searched on YouTube.")
    add_field(doc, "Handle Input", "Text input", "Competitors tab → add form",
        "Type the competitor's handle (without the @ symbol).",
        "Type 'mrbeast' (not '@mrbeast').")
    add_field(doc, "Add", "Button", "Competitors tab → add form",
        "Adds the competitor to your tracked list.",
        "Click Add → competitor appears in 'Tracked competitors' list.")
    add_field(doc, "Check now", "Button", "Competitors tab → top right",
        "Fetches latest posts from all tracked competitors. Takes 30-60 seconds per competitor.",
        "Click → 'Checking…' → 'Checked 3 competitors, found 8 new posts, 2 viral alerts'.",
        "Also auto-runs daily at 7am via cron.")
    add_field(doc, "Remove", "Button", "Competitors tab → competitor card",
        "Stops tracking a competitor. Their posts are deleted from the database.",
        "Click Remove → competitor is removed from your list.")
    add_field(doc, "Viral Alerts", "Display section", "Competitors tab → below tracked list",
        "Shows competitor posts that scored 70+ on the viral scale. Each includes AI suggestion for your own version.",
        "🚨 alert: 'MrBeast posted X (score 92) — Suggestion: Create a similar video about...'")
    add_field(doc, "View →", "Link", "Competitors tab → viral alert card",
        "Opens the competitor's original post in a new tab.",
        "Click → opens youtube.com/watch?v=xxx in new tab.")

    # ---- BRAND KIT ----
    doc.add_page_break()
    add_heading(doc, "Brand Kit Tab", level=2)

    add_field(doc, "Upload logo", "Button", "Brand Kit tab → Logo section",
        "Opens file picker for logo upload. PNG with transparency recommended.",
        "Click → select 'mylogo.png' → logo preview appears.",
        "300x300px is a good size. Avoid large files — they slow down video processing.")
    add_field(doc, "Replace logo", "Button", "Brand Kit tab → Logo section (when logo exists)",
        "Upload a new logo to replace the current one.",
        "Click → select new file → old logo is replaced.")
    add_field(doc, "Brand Name", "Text input", "Brand Kit tab → Brand Identity",
        "Your brand or channel name. Used in AI-generated content.",
        "'Puppy Adventures' or 'Fit with Sarah'.")
    add_field(doc, "Primary Color / Secondary / Accent", "Color pickers", "Brand Kit tab → Brand Identity",
        "Your brand's color palette in hex. Used by AI when generating thumbnails and assets.",
        "Primary: #FF6B35 (orange). Secondary: #FFFFFF (white). Accent: #000000 (black).")
    add_field(doc, "Font Family", "Dropdown", "Brand Kit tab → Brand Identity",
        "Font used for burned captions and text overlays.",
        "Select 'Inter' → captions use Inter font.",
        "Impact is great for thumbnail-style text. Inter/Montserrat for modern clean look.")
    add_field(doc, "Watermark Position", "Dropdown", "Brand Kit tab → Brand Identity",
        "Where your logo appears on videos. Same as the upload setting but applies globally.",
        "Select 'Bottom Right' → all future videos have logo in bottom right.")
    add_field(doc, "Watermark Size", "Slider (5-40%)", "Brand Kit tab → Brand Identity",
        "How large the watermark is relative to video width.",
        "15% (default) = logo takes up 15% of video width.",
        "Test different sizes. Too small = invisible. Too large = distracting.")
    add_field(doc, "Save Brand Kit", "Button", "Brand Kit tab → bottom",
        "Saves all brand kit settings. Applies to future video uploads.",
        "Click → 'Brand kit saved' toast appears.")
    add_field(doc, "Delete", "Button", "Brand Kit tab → bottom",
        "Deletes the entire brand kit. All settings reset to defaults.",
        "Click → 'Delete entire brand kit?' → click OK.")

    # ---- ASSETS ----
    doc.add_page_break()
    add_heading(doc, "Assets Tab", level=2)

    add_field(doc, "Watermark (PNG) — File Input", "File input", "Assets tab",
        "Upload a PNG logo to use as a watermark on videos. Appears in the Upload tab dropdown.",
        "Select 'logo.png' → uploaded → appears in Upload → Watermark dropdown.")
    add_field(doc, "Intro Clip (MP4) — File Input", "File input", "Assets tab",
        "Upload a video clip to prepend to videos. Appears in the Upload tab dropdown.",
        "Select 'intro.mp4' → uploaded → appears in Upload → Intro Clip dropdown.")
    add_field(doc, "Outro Clip (MP4) — File Input", "File input", "Assets tab",
        "Upload a video clip to append to videos. Appears in the Upload tab dropdown.",
        "Select 'outro.mp4' → uploaded → appears in Upload → Outro Clip dropdown.")
    add_field(doc, "Background Music (MP3) — File Input", "File input", "Assets tab",
        "Upload a music file to mix under videos. Appears in the Upload tab dropdown.",
        "Select 'upbeat.mp3' → uploaded → appears in Upload → Background Music dropdown.")
    add_field(doc, "Assets Table", "Display", "Assets tab → below upload forms",
        "Lists all uploaded assets with name, type, size, upload date, and delete button.",
        "Shows: 'logo.png · watermark · 45 KB · 6/25/2026 · Delete'.")
    add_field(doc, "Delete", "Button", "Assets tab → asset row",
        "Deletes an asset. It will no longer be available in the Upload dropdown.",
        "Click Delete → asset is removed.")

    # ---- SETTINGS ----
    doc.add_page_break()
    add_heading(doc, "Settings Tab", level=2)

    add_field(doc, "Brand Handle", "Text input", "Settings tab",
        "Your social media handle. Used in AI-generated captions and taglines.",
        "'@yourdog' or '@fitwithsarah'.",
        "Include the @ symbol. This appears at the end of AI-generated captions.")
    add_field(doc, "Content Niche", "Text input", "Settings tab",
        "What kind of content you make. AI uses this for ALL content generation — titles, captions, ideas, trends, insights.",
        "'dog / pet content' or 'fitness and workouts' or 'cooking and recipes'.",
        "Be specific. 'dog content' is okay. 'funny dog videos for TikTok' is better. This single field affects everything.")
    add_field(doc, "Save Settings", "Button", "Settings tab → bottom",
        "Saves your settings. Immediately affects all AI-generated content.",
        "Click → 'Settings saved' toast appears.")

    # ---- AI AGENT ----
    doc.add_page_break()
    add_heading(doc, "AI Agent Chat", level=2)
    add_para(doc, "The floating orange button in the bottom-right corner. Click to open the chat sidebar.", italic=True, size=10)

    add_field(doc, "Chat Input", "Textarea", "Agent chat → bottom",
        "Type any question or command. Press Enter to send (Shift+Enter for new line).",
        "Type 'What should I post this week?' → press Enter → agent responds.",
        "The agent can call up to 33 tools. It sees your data and can take actions.")
    add_field(doc, "Send Button", "Button", "Agent chat → bottom right",
        "Sends your message. Disabled while agent is thinking.",
        "Click → message is sent → agent processes → response appears.")
    add_field(doc, "Suggested Questions", "Buttons", "Agent chat → empty state",
        "4 pre-written questions to get you started. Click any to fill the input.",
        "Click 'Show me my top performing videos' → input is filled → press Enter to send.")
    add_field(doc, "Tool Call Indicators", "Display (chips)", "Agent chat → response",
        "Shows which tools the agent called. Click 'view result' to see the data returned.",
        "Shows 'used tool: list_videos' with expandable details.")
    add_field(doc, "New Conversation", "Button", "Agent chat → header",
        "Starts a fresh conversation. Previous conversation is saved and can be referenced later.",
        "Click → empty chat appears → type a new question.")

    add_para(doc, "")
    add_para(doc, "Ready-to-use Agent Commands:", bold=True, size=12, space_after=4)
    commands = [
        "What should I post this week?",
        "Give me 3 content ideas",
        "What's trending in my niche?",
        "Schedule my ready videos at optimal times",
        "What videos do I have?",
        "Show me my top performing videos",
        "Why did my last video underperform?",
        "Generate a thumbnail for my latest video",
        "Repurpose my latest video into 5 clips",
        "Translate my latest video to Spanish and French",
        "Generate 5 hooks for my latest video",
        "What are my competitors up to?",
        "Fetch new comments and suggest replies",
        "What is my brand name and primary color?",
        "Add @competitor_handle on YouTube as a competitor",
    ]
    for cmd in commands:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"• \"{cmd}\"")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xE8, 0x6A, 0x00)

    # ============ SAVE ============
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")

if __name__ == "__main__":
    build_document()
